"""Word file manipulation tool for Gemini Office Agent.

This module provides the WordTool class for reading, creating, and modifying
Word files (.docx) using the python-docx library.
"""

import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError

from src.exceptions import CorruptedFileError
from src.logging_config import get_logger, log_file_access_error

logger = get_logger(__name__)


class WordTool:
    """Tool for manipulating Word files (.docx).
    
    Provides methods for reading, creating, updating, and adding content
    to Word documents using python-docx.
    """
    
    def __init__(self, gemini_client=None):
        """Initialize WordTool with optional Gemini client for AI-powered operations.
        
        Args:
            gemini_client: Optional GeminiClient instance for text improvement operations
        """
        self.gemini_client = gemini_client

    def read_word(self, file_path: str) -> str:
        """Lê documento Word e retorna texto completo.
        
        Args:
            file_path: Path to the Word file to read
            
        Returns:
            Complete text content from the document (all paragraphs concatenated)
                
        Raises:
            FileNotFoundError: If the file does not exist
            CorruptedFileError: If the file is corrupted or cannot be read
        """
        logger.info(f"Reading Word file: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract all paragraph text
            paragraphs = [para.text for para in doc.paragraphs]
            full_text = "\n".join(paragraphs)
            
            logger.info(f"Successfully read Word file with {len(paragraphs)} paragraphs")
            return full_text
            
        except PackageNotFoundError as e:
            logger.error(f"Corrupted Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error reading Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path} - {str(e)}") from e

    def create_word(self, file_path: str, content: str) -> None:
        """Cria novo documento Word com conteúdo.
        
        Args:
            file_path: Path where the Word file should be created
            content: Text content to add to the document (can contain newlines for multiple paragraphs)
                    Can be a string or a list of strings (one per paragraph)
            
        Raises:
            IOError: If the file cannot be created
        """
        logger.info(f"Creating Word file: {file_path}")
        
        Path(file_path).parent.mkdir(parents=True, exist_ok=True)
        
        try:
            doc = Document()
            
            # Handle both string and list content
            if isinstance(content, list):
                lines = content
            elif isinstance(content, str):
                lines = content.split("\n")
            else:
                # Fallback: convert to string
                lines = [str(content)]
            
            # Add each line as a paragraph
            for line in lines:
                doc.add_paragraph(line)
            
            # Save the document
            doc.save(file_path)
            
            logger.info(f"Successfully created Word file with {len(lines)} paragraphs")
            
        except Exception as e:
            logger.error(f"Error creating Word file: {file_path} - {str(e)}")
            raise IOError(f"Failed to create Word file: {file_path} - {str(e)}") from e

    def add_paragraph(self, file_path: str, text: str) -> None:
        """Adiciona parágrafo ao documento.
        
        Args:
            file_path: Path to the Word file
            text: Text content for the new paragraph
            
        Raises:
            FileNotFoundError: If the file does not exist
            CorruptedFileError: If the file is corrupted
        """
        logger.info(f"Adding paragraph to {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Add new paragraph
            doc.add_paragraph(text)
            
            # Save the document
            doc.save(file_path)
            
            logger.info(f"Successfully added paragraph to document")
            
        except PackageNotFoundError as e:
            logger.error(f"Corrupted Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding paragraph to Word file: {file_path} - {str(e)}")
            raise IOError(f"Failed to add paragraph to Word file: {file_path} - {str(e)}") from e

    def update_paragraph(self, file_path: str, index: int, new_text: str) -> None:
        """Atualiza parágrafo específico.
        
        Args:
            file_path: Path to the Word file
            index: Index of the paragraph to update (0-based)
            new_text: New text content for the paragraph
            
        Raises:
            FileNotFoundError: If the file does not exist
            CorruptedFileError: If the file is corrupted
            IndexError: If the paragraph index is out of range
        """
        logger.info(f"Updating paragraph {index} in {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Check if index is valid
            if index < 0 or index >= len(doc.paragraphs):
                logger.error(f"Paragraph index {index} out of range (document has {len(doc.paragraphs)} paragraphs)")
                raise IndexError(f"Paragraph index {index} out of range (document has {len(doc.paragraphs)} paragraphs)")
            
            # Update paragraph text by clearing runs and adding new text
            paragraph = doc.paragraphs[index]
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            # Add new text
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
            
            # Save the document
            doc.save(file_path)
            
            logger.info(f"Successfully updated paragraph {index}")
            
        except PackageNotFoundError as e:
            logger.error(f"Corrupted Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except IndexError:
            raise
        except Exception as e:
            logger.error(f"Error updating paragraph in Word file: {file_path} - {str(e)}")
            raise IOError(f"Failed to update paragraph in Word file: {file_path} - {str(e)}") from e

    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """Extrai informações de tabelas do documento.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            List of tables, where each table is a list of rows,
            and each row is a list of cell values (as strings)
                
        Raises:
            FileNotFoundError: If the file does not exist
            CorruptedFileError: If the file is corrupted or cannot be read
        """
        logger.info(f"Extracting tables from Word file: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract all tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            logger.info(f"Successfully extracted {len(tables_data)} tables from document")
            return tables_data
            
        except PackageNotFoundError as e:
            logger.error(f"Corrupted Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error extracting tables from Word file: {file_path} - {str(e)}")
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path} - {str(e)}") from e

    def add_heading(self, file_path: str, text: str, level: int = 1) -> None:
        """Adiciona um título/heading ao documento.
        
        Args:
            file_path: Path to the Word file
            text: Text content for the heading
            level: Heading level (0=Title, 1=Heading1, 2=Heading2, etc., max 9)
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the level is invalid
        """
        # Ensure level is int (Gemini API may send strings)
        level = int(level)

        logger.info(f"Adding heading (level {level}) to {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        if level < 0 or level > 9:
            raise ValueError(f"Heading level must be between 0 and 9, got {level}")
        
        try:
            doc = Document(file_path)
            doc.add_heading(text, level=level)
            doc.save(file_path)
            logger.info(f"Successfully added heading level {level}: '{text[:50]}'")
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error adding heading: {file_path} - {str(e)}")
            raise IOError(f"Failed to add heading: {file_path} - {str(e)}") from e

    def add_table(self, file_path: str, headers: List[str], rows: List[List[str]],
                  style: str = "Table Grid") -> None:
        """Adiciona uma tabela ao documento.
        
        Args:
            file_path: Path to the Word file
            headers: List of header column names
            rows: List of rows, each row is a list of cell values
            style: Table style name (default: "Table Grid")
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Adding table ({len(headers)} cols, {len(rows)} rows) to {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        # Validate table dimensions to prevent corrupted documents
        if not headers or len(headers) == 0:
            raise ValueError("A tabela precisa ter ao menos 1 coluna (headers não pode ser vazio)")
        
        try:
            doc = Document(file_path)
            
            total_rows = len(rows) + 1  # +1 for header
            total_cols = len(headers)
            table = doc.add_table(rows=total_rows, cols=total_cols)
            
            try:
                table.style = style
            except KeyError:
                logger.warning(f"Table style '{style}' not found, using default")
            
            # Header row
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data rows
            for row_idx, row_data in enumerate(rows):
                for col_idx, value in enumerate(row_data):
                    if col_idx < total_cols:
                        table.rows[row_idx + 1].cells[col_idx].text = str(value)
            
            doc.save(file_path)
            logger.info(f"Successfully added table with {len(rows)} data rows")
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding table: {file_path} - {str(e)}")
            raise IOError(f"Failed to add table: {file_path} - {str(e)}") from e

    def delete_paragraph(self, file_path: str, index: int) -> None:
        """Remove um parágrafo do documento pelo índice.
        
        Args:
            file_path: Path to the Word file
            index: Index of the paragraph to delete (0-based)
            
        Raises:
            FileNotFoundError: If the file does not exist
            IndexError: If the paragraph index is out of range
        """
        logger.info(f"Deleting paragraph {index} from {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            if index < 0 or index >= len(doc.paragraphs):
                raise IndexError(
                    f"Paragraph index {index} out of range "
                    f"(document has {len(doc.paragraphs)} paragraphs)"
                )
            
            paragraph = doc.paragraphs[index]
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            
            doc.save(file_path)
            logger.info(f"Successfully deleted paragraph {index}")
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except IndexError:
            raise
        except Exception as e:
            logger.error(f"Error deleting paragraph: {file_path} - {str(e)}")
            raise IOError(f"Failed to delete paragraph: {file_path} - {str(e)}") from e

    def replace_text(self, file_path: str, old_text: str, new_text: str) -> int:
        """Busca e substitui texto em todo o documento (parágrafos + tabelas).
        
        Args:
            file_path: Path to the Word file
            old_text: Text to search for
            new_text: Replacement text
            
        Returns:
            Number of replacements made
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Replacing '{old_text[:30]}' with '{new_text[:30]}' in {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            count = 0
            
            # Substituir em parágrafos
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
            
            # Substituir em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if old_text in paragraph.text:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        count += 1
            
            doc.save(file_path)
            logger.info(f"Successfully replaced {count} occurrence(s)")
            return count
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error replacing text: {file_path} - {str(e)}")
            raise IOError(f"Failed to replace text: {file_path} - {str(e)}") from e

    def format_paragraph(self, file_path: str, index: int, formatting: Dict[str, Any]) -> None:
        """Aplica formatação a um parágrafo específico.

        Args:
            file_path: Path to the Word file
            index: Index of the paragraph (0-based)
            formatting: Dict com configurações de formatação:
                Texto básico:
                - bold: bool — negrito
                - italic: bool — itálico
                - underline: bool — sublinhado
                - font_size: int/float — tamanho em pontos (ex: 12)
                - font_color: str hex sem '#' (ex: "FF0000")
                - font_name: str (ex: "Arial", "Calibri", "Times New Roman")
                - alignment: str "left"|"center"|"right"|"justify"

                Formatação adicional:
                - strikethrough: bool — texto tachado (riscado)
                - highlight: str — cor de destaque: "yellow", "green", "cyan", "magenta",
                                   "blue", "red", "dark_blue", "dark_cyan", "dark_green",
                                   "dark_magenta", "dark_red", "dark_yellow", "dark_gray",
                                   "light_gray", "black", "white"
                - superscript: bool — texto sobrescrito (ex: x²)
                - subscript: bool — texto subscrito (ex: H₂O)
                - all_caps: bool — todas as letras em maiúsculas
                - small_caps: bool — versaletes (maiúsculas em tamanho reduzido)
                - space_before: float — espaço antes do parágrafo em pontos
                - space_after: float — espaço depois do parágrafo em pontos
                - line_spacing: float — espaçamento entre linhas em pontos

        Raises:
            FileNotFoundError: If the file does not exist
            IndexError: If paragraph index is out of range
        """
        # Ensure index is int (Gemini API may send strings)
        if index is not None:
            index = int(index)

        logger.info(f"Formatting paragraph {index} in {file_path}")

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.enum.text import WD_COLOR_INDEX
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Pt as _Pt

            # Ensure numeric formatting values are proper types
            for num_key in ('font_size', 'space_before', 'space_after', 'line_spacing'):
                if num_key in formatting and formatting[num_key] is not None:
                    formatting[num_key] = float(formatting[num_key])

            # Highlight color name → WD_COLOR_INDEX mapping
            # Valid WD_COLOR_INDEX values: AUTO, BLACK, BLUE, BRIGHT_GREEN,
            # DARK_BLUE, DARK_RED, DARK_YELLOW, GRAY_25, GRAY_50, GREEN,
            # PINK, RED, TEAL, TURQUOISE, VIOLET, WHITE, YELLOW
            HIGHLIGHT_MAP = {
                'yellow':       WD_COLOR_INDEX.YELLOW,
                'green':        WD_COLOR_INDEX.BRIGHT_GREEN,
                'cyan':         WD_COLOR_INDEX.TURQUOISE,
                'magenta':      WD_COLOR_INDEX.PINK,
                'blue':         WD_COLOR_INDEX.BLUE,
                'red':          WD_COLOR_INDEX.RED,
                'dark_blue':    WD_COLOR_INDEX.DARK_BLUE,
                'dark_cyan':    WD_COLOR_INDEX.TEAL,
                'dark_green':   WD_COLOR_INDEX.GREEN,
                'dark_magenta': WD_COLOR_INDEX.VIOLET,
                'dark_red':     WD_COLOR_INDEX.DARK_RED,
                'dark_yellow':  WD_COLOR_INDEX.DARK_YELLOW,
                'dark_gray':    WD_COLOR_INDEX.GRAY_50,
                'light_gray':   WD_COLOR_INDEX.GRAY_25,
                'black':        WD_COLOR_INDEX.BLACK,
                'white':        WD_COLOR_INDEX.WHITE,
                # Additional aliases for convenience
                'pink':         WD_COLOR_INDEX.PINK,
                'turquoise':    WD_COLOR_INDEX.TURQUOISE,
                'teal':         WD_COLOR_INDEX.TEAL,
                'violet':       WD_COLOR_INDEX.VIOLET,
                'bright_green': WD_COLOR_INDEX.BRIGHT_GREEN,
                'gray_25':      WD_COLOR_INDEX.GRAY_25,
                'gray_50':      WD_COLOR_INDEX.GRAY_50,
                'gray':         WD_COLOR_INDEX.GRAY_50,
                'auto':         WD_COLOR_INDEX.AUTO,
            }

            doc = Document(file_path)

            # Default to last paragraph if index is None
            if index is None:
                if len(doc.paragraphs) == 0:
                    raise IndexError("Document has no paragraphs to format")
                index = len(doc.paragraphs) - 1
                logger.info(f"Index not provided, defaulting to last paragraph (index={index})")

            if index < 0 or index >= len(doc.paragraphs):
                raise IndexError(
                    f"Paragraph index {index} out of range "
                    f"(document has {len(doc.paragraphs)} paragraphs)"
                )

            paragraph = doc.paragraphs[index]

            # --- Paragraph-level formatting ---
            align_map = {
                'left':    WD_ALIGN_PARAGRAPH.LEFT,
                'center':  WD_ALIGN_PARAGRAPH.CENTER,
                'right':   WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            alignment = formatting.get('alignment')
            if alignment and alignment.lower() in align_map:
                paragraph.alignment = align_map[alignment.lower()]

            # Paragraph spacing
            if 'space_before' in formatting:
                paragraph.paragraph_format.space_before = _Pt(formatting['space_before'])
            if 'space_after' in formatting:
                paragraph.paragraph_format.space_after = _Pt(formatting['space_after'])
            if 'line_spacing' in formatting:
                paragraph.paragraph_format.line_spacing = _Pt(formatting['line_spacing'])

            # --- Run-level formatting ---
            for run in paragraph.runs:
                # Basic
                if 'bold' in formatting:
                    run.bold = formatting['bold']
                if 'italic' in formatting:
                    run.italic = formatting['italic']
                if 'underline' in formatting:
                    run.underline = formatting['underline']
                if 'font_size' in formatting:
                    run.font.size = Pt(formatting['font_size'])
                if 'font_name' in formatting:
                    run.font.name = formatting['font_name']
                if 'font_color' in formatting:
                    hex_color = formatting['font_color'].lstrip('#')
                    run.font.color.rgb = RGBColor(
                        int(hex_color[0:2], 16),
                        int(hex_color[2:4], 16),
                        int(hex_color[4:6], 16)
                    )

                # Strikethrough (tachado)
                if 'strikethrough' in formatting:
                    run.font.strike = formatting['strikethrough']

                # Highlight (marca-texto)
                if 'highlight' in formatting:
                    color_key = formatting['highlight'].lower()
                    if color_key in HIGHLIGHT_MAP:
                        run.font.highlight_color = HIGHLIGHT_MAP[color_key]

                # Superscript / Subscript
                if 'superscript' in formatting:
                    run.font.superscript = formatting['superscript']
                if 'subscript' in formatting:
                    run.font.subscript = formatting['subscript']

                # All caps / Small caps
                if 'all_caps' in formatting:
                    run.font.all_caps = formatting['all_caps']
                if 'small_caps' in formatting:
                    run.font.small_caps = formatting['small_caps']

            doc.save(file_path)
            logger.info(f"Successfully formatted paragraph {index}")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except IndexError:
            raise
        except Exception as e:
            logger.error(f"Error formatting paragraph: {file_path} - {str(e)}")
            raise IOError(f"Failed to format paragraph: {file_path} - {str(e)}") from e

    def add_list(self, file_path: str, items: List[str], ordered: bool = False) -> None:
        """Adiciona uma lista (bullet ou numerada) ao documento.
        
        Args:
            file_path: Path to the Word file
            items: List of text items
            ordered: True for numbered list, False for bullet list
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        list_type = "numbered" if ordered else "bullet"
        logger.info(f"Adding {list_type} list ({len(items)} items) to {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            style = 'List Number' if ordered else 'List Bullet'
            for item in items:
                doc.add_paragraph(item, style=style)
            
            doc.save(file_path)
            logger.info(f"Successfully added {list_type} list with {len(items)} items")
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding list: {file_path} - {str(e)}")
            raise IOError(f"Failed to add list: {file_path} - {str(e)}") from e

    def add_image(self, file_path: str, image_path: str, 
                  width: Optional[float] = None, height: Optional[float] = None,
                  alignment: str = "left", caption: Optional[str] = None) -> None:
        """Adiciona uma imagem ao documento Word.
        
        Suporta formatos: PNG, JPG, JPEG, GIF, BMP, TIFF, EMF, WMF.
        Ideal para fotos de currículo, logos, gráficos e ilustrações.
        
        Args:
            file_path: Path to the Word file
            image_path: Path to the image file to insert
            width: Optional width in inches (e.g., 2.0 for 2 inches)
            height: Optional height in inches (maintains aspect ratio if only one specified)
            alignment: Image alignment - 'left', 'center', or 'right'
            caption: Optional caption text below the image
            
        Raises:
            FileNotFoundError: If the Word file or image does not exist
            ValueError: If image format is not supported
            IOError: If the image cannot be added
            
        Examples:
            # Add profile photo for resume (2 inches wide, centered)
            add_image("resume.docx", "photo.jpg", width=2.0, alignment="center")
            
            # Add logo with caption
            add_image("report.docx", "logo.png", width=1.5, caption="Company Logo")
            
            # Add full-width image
            add_image("doc.docx", "chart.png", width=6.0, alignment="center")
        """
        logger.info(f"Adding image to {file_path}: {image_path}")
        
        # Validate Word file exists
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        # Validate image file exists
        img_path = Path(image_path)
        if not img_path.exists():
            log_file_access_error(logger, image_path, "imagem não encontrada")
            raise FileNotFoundError(f"Image file not found: {image_path}")
        
        # Validate image format
        supported_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.emf', '.wmf'}
        if img_path.suffix.lower() not in supported_formats:
            raise ValueError(
                f"Formato de imagem não suportado: {img_path.suffix}. "
                f"Formatos suportados: {', '.join(sorted(supported_formats))}"
            )
        
        try:
            doc = Document(file_path)
            
            # Create paragraph for the image
            paragraph = doc.add_paragraph()
            
            # Set alignment
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }
            paragraph.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
            
            # Add the image
            run = paragraph.add_run()
            
            # Determine dimensions
            if width and height:
                run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
            elif width:
                run.add_picture(str(image_path), width=Inches(width))
            elif height:
                run.add_picture(str(image_path), height=Inches(height))
            else:
                # Default: add with original size (may be large)
                run.add_picture(str(image_path))
            
            # Add caption if provided
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
                caption_run = caption_para.add_run(caption)
                caption_run.italic = True
                caption_run.font.size = Pt(10)
            
            doc.save(file_path)
            
            size_info = f"width={width}" if width else ("height={height}" if height else "original size")
            logger.info(f"Successfully added image ({size_info}, {alignment})")
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding image: {file_path} - {str(e)}")
            raise IOError(f"Failed to add image: {file_path} - {str(e)}") from e

    def add_image_at_position(self, file_path: str, image_path: str, 
                               paragraph_index: int,
                               width: Optional[float] = None, 
                               height: Optional[float] = None) -> None:
        """Insere uma imagem em uma posição específica do documento.
        
        A imagem é inserida ANTES do parágrafo especificado.
        
        Args:
            file_path: Path to the Word file
            image_path: Path to the image file
            paragraph_index: Index of paragraph before which to insert (0-based)
            width: Optional width in inches
            height: Optional height in inches
            
        Raises:
            FileNotFoundError: If files don't exist
            IndexError: If paragraph_index is out of range
        """
        logger.info(f"Adding image at position {paragraph_index} in {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        img_path = Path(image_path)
        if not img_path.exists():
            raise FileNotFoundError(f"Image file not found: {image_path}")
        
        try:
            doc = Document(file_path)
            
            if paragraph_index < 0 or paragraph_index > len(doc.paragraphs):
                raise IndexError(
                    f"Índice de parágrafo inválido: {paragraph_index}. "
                    f"Documento tem {len(doc.paragraphs)} parágrafos."
                )
            
            # Create new paragraph for image
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            
            if width and height:
                run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
            elif width:
                run.add_picture(str(image_path), width=Inches(width))
            elif height:
                run.add_picture(str(image_path), height=Inches(height))
            else:
                run.add_picture(str(image_path))
            
            # Move paragraph to correct position (insert before target paragraph)
            if paragraph_index < len(doc.paragraphs) - 1:
                target_para = doc.paragraphs[paragraph_index]
                target_para._element.addprevious(new_para._element)
            
            doc.save(file_path)
            logger.info(f"Successfully added image at position {paragraph_index}")
            
        except Exception as e:
            logger.error(f"Error adding image at position: {str(e)}")
            raise

    def add_hyperlink(self, file_path: str, text: str, url: str,
                      bold: bool = False, italic: bool = False,
                      font_size: Optional[float] = None,
                      color: str = "0563C1") -> None:
        """Adiciona um parágrafo contendo um hyperlink clicável ao documento.

        Usa manipulação XML direta (python-docx não possui API nativa para links).
        Ideal para currículos (LinkedIn, portfólio, email) e relatórios.

        Args:
            file_path: Path to the Word file
            text: Display text for the hyperlink
            url: Target URL (e.g., 'https://linkedin.com/in/usuario' or 'mailto:email@email.com')
            bold: Whether the link text should be bold
            italic: Whether the link text should be italic
            font_size: Optional font size in points
            color: Hex color for the link text without '#' (default: Word blue '0563C1')

        Raises:
            FileNotFoundError: If the Word file does not exist
            ValueError: If url or text is empty
            IOError: If the hyperlink cannot be added

        Examples:
            # LinkedIn link on a resume
            add_hyperlink("curriculo.docx", "linkedin.com/in/joao", "https://linkedin.com/in/joao")

            # Email link
            add_hyperlink("curriculo.docx", "joao@email.com", "mailto:joao@email.com")

            # Portfolio link, bold
            add_hyperlink("doc.docx", "Ver Portfólio", "https://portfolio.com", bold=True)
        """
        logger.info(f"Adding hyperlink to {file_path}: '{text}' -> {url}")

        if not text or not text.strip():
            raise ValueError("O texto do hyperlink não pode ser vazio")
        if not url or not url.strip():
            raise ValueError("A URL do hyperlink não pode ser vazia")

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import re as _re

            doc = Document(file_path)

            # Add a new paragraph to hold the hyperlink
            paragraph = doc.add_paragraph()

            # Register the relationship for the hyperlink target
            part = doc.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

            # Build <w:hyperlink r:id="..." w:history="1"> XML element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.set(qn('w:history'), '1')

            # Build <w:r> run inside the hyperlink
            run_elem = OxmlElement('w:r')

            # Build <w:rPr> run properties
            rpr = OxmlElement('w:rPr')

            # Apply hyperlink character style
            r_style = OxmlElement('w:rStyle')
            r_style.set(qn('w:val'), 'Hyperlink')
            rpr.append(r_style)

            # Font color
            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), color.upper().lstrip('#'))
            rpr.append(color_elem)

            # Underline (standard for hyperlinks)
            u_elem = OxmlElement('w:u')
            u_elem.set(qn('w:val'), 'single')
            rpr.append(u_elem)

            # Bold
            if bold:
                b_elem = OxmlElement('w:b')
                rpr.append(b_elem)

            # Italic
            if italic:
                i_elem = OxmlElement('w:i')
                rpr.append(i_elem)

            # Font size (half-points in OOXML)
            if font_size:
                sz_elem = OxmlElement('w:sz')
                sz_elem.set(qn('w:val'), str(int(font_size * 2)))
                rpr.append(sz_elem)
                sz_cs = OxmlElement('w:szCs')
                sz_cs.set(qn('w:val'), str(int(font_size * 2)))
                rpr.append(sz_cs)

            run_elem.append(rpr)

            # Text content <w:t>
            t_elem = OxmlElement('w:t')
            t_elem.text = text
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            run_elem.append(t_elem)

            hyperlink.append(run_elem)
            paragraph._p.append(hyperlink)

            doc.save(file_path)
            logger.info(f"Successfully added hyperlink: '{text}' -> {url}")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding hyperlink: {file_path} - {str(e)}")
            raise IOError(f"Failed to add hyperlink: {file_path} - {str(e)}") from e

    def add_hyperlink_to_paragraph(self, file_path: str, paragraph_index: int,
                                    text: str, url: str,
                                    bold: bool = False, italic: bool = False,
                                    color: str = "0563C1") -> None:
        """Insere um hyperlink inline dentro de um parágrafo existente.

        O hyperlink é adicionado ao FINAL do parágrafo especificado.

        Args:
            file_path: Path to the Word file
            paragraph_index: Index of the existing paragraph (0-based)
            text: Display text for the hyperlink
            url: Target URL
            bold: Whether the link text should be bold
            italic: Whether the link text should be italic
            color: Hex color for the link text without '#' (default: '0563C1')

        Raises:
            FileNotFoundError: If the Word file does not exist
            IndexError: If paragraph_index is out of range
            IOError: If the hyperlink cannot be added
        """
        logger.info(f"Adding inline hyperlink at paragraph {paragraph_index} in {file_path}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document(file_path)

            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                raise IndexError(
                    f"Índice de parágrafo inválido: {paragraph_index}. "
                    f"Documento tem {len(doc.paragraphs)} parágrafos."
                )

            paragraph = doc.paragraphs[paragraph_index]
            part = doc.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            hyperlink.set(qn('w:history'), '1')

            run_elem = OxmlElement('w:r')
            rpr = OxmlElement('w:rPr')

            r_style = OxmlElement('w:rStyle')
            r_style.set(qn('w:val'), 'Hyperlink')
            rpr.append(r_style)

            color_elem = OxmlElement('w:color')
            color_elem.set(qn('w:val'), color.upper().lstrip('#'))
            rpr.append(color_elem)

            u_elem = OxmlElement('w:u')
            u_elem.set(qn('w:val'), 'single')
            rpr.append(u_elem)

            if bold:
                rpr.append(OxmlElement('w:b'))
            if italic:
                rpr.append(OxmlElement('w:i'))

            run_elem.append(rpr)

            t_elem = OxmlElement('w:t')
            t_elem.text = text
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            run_elem.append(t_elem)

            hyperlink.append(run_elem)
            paragraph._p.append(hyperlink)

            doc.save(file_path)
            logger.info(f"Successfully added inline hyperlink at paragraph {paragraph_index}")

        except Exception as e:
            logger.error(f"Error adding inline hyperlink: {str(e)}")
            raise

    # ==================== HEADER AND FOOTER METHODS ====================

    def add_header(self, file_path: str, text: str,
                   alignment: str = "center",
                   bold: bool = False,
                   italic: bool = False,
                   font_size: Optional[float] = None,
                   font_name: Optional[str] = None,
                   include_page_number: bool = False,
                   page_number_position: str = "right",
                   include_total_pages: bool = False,
                   use_tab_stops: bool = False) -> None:
        """Adiciona ou substitui o cabeçalho do documento Word.

        O cabeçalho aparece no topo de todas as páginas.
        Suporta texto livre, formatação e número de página opcional.

        Args:
            file_path: Path to the Word file
            text: Header text content
            alignment: Text alignment - 'left', 'center', or 'right'
            bold: Whether the header text should be bold
            italic: Whether the header text should be italic
            font_size: Font size in points (e.g., 10.0)
            font_name: Font name (e.g., 'Arial', 'Calibri')
            include_page_number: Whether to add a page number field
            page_number_position: Where to place page number - 'left', 'center', or 'right'
            include_total_pages: If True, adds ' de X' after page number (e.g., 'Página 1 de 5')
            use_tab_stops: If True, configures tab stops to properly align text left and page number right

        Raises:
            FileNotFoundError: If the Word file does not exist
            IOError: If the header cannot be added

        Examples:
            # Simple centered header
            add_header("relatorio.docx", "Relatório Anual 2025", alignment="center")

            # Header with page number on the right
            add_header("doc.docx", "Empresa XYZ", include_page_number=True, page_number_position="right")

            # Bold header with custom font
            add_header("doc.docx", "Confidencial", bold=True, font_name="Arial", font_size=10)
            
            # Header with text left and "Page X of Y" right (same line)
            add_header("doc.docx", "Company Report", alignment="left", 
                      include_page_number=True, page_number_position="right",
                      include_total_pages=True, use_tab_stops=True)
        """
        logger.info(f"Adding header to {file_path}: '{text}'")

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

            doc = Document(file_path)

            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }

            for section in doc.sections:
                section.header.is_linked_to_previous = False
                header = section.header

                # Clear existing header content
                for para in header.paragraphs:
                    for run in para.runs:
                        run.text = ''

                # Use first paragraph or add one
                if header.paragraphs:
                    para = header.paragraphs[0]
                    # Remove all child run elements so paragraph is fully empty
                    for run in list(para.runs):
                        run._r.getparent().remove(run._r)
                else:
                    para = header.add_paragraph()

                para.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)

                # Configure tab stops if needed (for text left + page number right)
                if use_tab_stops and alignment.lower() == 'left' and page_number_position.lower() == 'right':
                    # Add a right-aligned tab stop at the right margin
                    # Page width minus margins (in twips: 1 inch = 1440 twips)
                    page_width = section.page_width
                    right_margin = section.right_margin
                    left_margin = section.left_margin
                    tab_position = page_width - right_margin - left_margin
                    
                    tab_stops = para.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)

                # Add main text run
                if text:
                    run = para.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    if font_size:
                        run.font.size = Pt(font_size)
                    if font_name:
                        run.font.name = font_name

                # Add page number field if requested
                if include_page_number:
                    if use_tab_stops and alignment.lower() == 'left' and page_number_position.lower() == 'right':
                        # Use tab to push page number to the right
                        para.add_run('\t')
                    elif page_number_position != alignment:
                        # Add a tab to separate text from page number (legacy behavior)
                        para.add_run('\t')

                    self._add_page_number_field(para)
                    
                    # Add total pages if requested: " de X"
                    if include_total_pages:
                        run_sep = para.add_run(' de ')
                        if font_size:
                            run_sep.font.size = Pt(font_size)
                        if font_name:
                            run_sep.font.name = font_name
                        self._add_num_pages_field(para)

            doc.save(file_path)
            logger.info(f"Successfully added header: '{text}'")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding header: {file_path} - {str(e)}")
            raise IOError(f"Failed to add header: {file_path} - {str(e)}") from e

    def add_footer(self, file_path: str, text: str = "",
                   alignment: str = "center",
                   bold: bool = False,
                   italic: bool = False,
                   font_size: Optional[float] = None,
                   font_name: Optional[str] = None,
                   include_page_number: bool = True,
                   page_number_position: str = "center",
                   include_total_pages: bool = False) -> None:
        """Adiciona ou substitui o rodapé do documento Word.

        O rodapé aparece na base de todas as páginas.
        Suporta texto livre, número de página e total de páginas (ex: 'Página 1 de 5').

        Args:
            file_path: Path to the Word file
            text: Footer text content (can be empty if only page number is desired)
            alignment: Text alignment - 'left', 'center', or 'right'
            bold: Whether the footer text should be bold
            italic: Whether the footer text should be italic
            font_size: Font size in points
            font_name: Font name
            include_page_number: Whether to add a page number field (default: True)
            page_number_position: Where to place page number - 'left', 'center', or 'right'
            include_total_pages: If True, adds ' de X' after page number (e.g., 'Página 1 de 5')

        Raises:
            FileNotFoundError: If the Word file does not exist
            IOError: If the footer cannot be added

        Examples:
            # Simple page number footer
            add_footer("relatorio.docx", "", include_page_number=True, page_number_position="center")

            # Footer with company name and page number
            add_footer("doc.docx", "Empresa XYZ", include_page_number=True, page_number_position="right")

            # Footer with "Página X de Y" format
            add_footer("doc.docx", "Página ", include_page_number=True, include_total_pages=True)
        """
        logger.info(f"Adding footer to {file_path}: '{text}'")

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = Document(file_path)

            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }

            for section in doc.sections:
                section.footer.is_linked_to_previous = False
                footer = section.footer

                # Use first paragraph or add one
                if footer.paragraphs:
                    para = footer.paragraphs[0]
                    para.clear()
                else:
                    para = footer.add_paragraph()

                para.alignment = align_map.get(page_number_position.lower(), WD_ALIGN_PARAGRAPH.CENTER)

                # Add text if provided
                if text:
                    run = para.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    if font_size:
                        run.font.size = Pt(font_size)
                    if font_name:
                        run.font.name = font_name

                # Add page number field
                if include_page_number:
                    self._add_page_number_field(para)

                    # Add total pages if requested: " de X"
                    if include_total_pages:
                        run_sep = para.add_run(' de ')
                        if font_size:
                            run_sep.font.size = Pt(font_size)
                        self._add_num_pages_field(para)

            doc.save(file_path)
            logger.info(f"Successfully added footer")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding footer: {file_path} - {str(e)}")
            raise IOError(f"Failed to add footer: {file_path} - {str(e)}") from e

    def remove_header(self, file_path: str) -> None:
        """Remove o cabeçalho de todas as seções do documento.

        Args:
            file_path: Path to the Word file

        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Removing header from {file_path}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = Document(file_path)
            for section in doc.sections:
                header = section.header
                for para in header.paragraphs:
                    para.clear()
            doc.save(file_path)
            logger.info("Header removed successfully")
        except Exception as e:
            logger.error(f"Error removing header: {str(e)}")
            raise

    def remove_footer(self, file_path: str) -> None:
        """Remove o rodapé de todas as seções do documento.

        Args:
            file_path: Path to the Word file

        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Removing footer from {file_path}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = Document(file_path)
            for section in doc.sections:
                footer = section.footer
                for para in footer.paragraphs:
                    para.clear()
            doc.save(file_path)
            logger.info("Footer removed successfully")
        except Exception as e:
            logger.error(f"Error removing footer: {str(e)}")
            raise

    def _add_page_number_field(self, paragraph) -> None:
        """Insere campo de número de página atual (PAGE) no parágrafo.

        Usa XML OOXML field codes: fldChar begin + instrText PAGE + fldChar end.

        Args:
            paragraph: python-docx Paragraph object
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        run = paragraph.add_run()
        run_elem = run._r

        # <w:fldChar w:fldCharType="begin"/>
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run_elem.append(fld_begin)

        # <w:instrText xml:space="preserve"> PAGE </w:instrText>
        run2 = paragraph.add_run()
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = ' PAGE '
        run2._r.append(instr)

        # <w:fldChar w:fldCharType="end"/>
        run3 = paragraph.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_end)

    def _add_num_pages_field(self, paragraph) -> None:
        """Insere campo de total de páginas (NUMPAGES) no parágrafo.

        Args:
            paragraph: python-docx Paragraph object
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        run = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        run2 = paragraph.add_run()
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = ' NUMPAGES '
        run2._r.append(instr)

        run3 = paragraph.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_end)

    # ==================== PAGE LAYOUT METHODS ====================

    def set_page_margins(self, file_path: str,
                         top: float = 2.5, bottom: float = 2.5,
                         left: float = 3.0, right: float = 2.0,
                         unit: str = 'cm') -> None:
        """Define as margens da página para todas as seções do documento.

        Args:
            file_path: Path to the Word file
            top: Top margin (default: 2.5 cm)
            bottom: Bottom margin (default: 2.5 cm)
            left: Left margin (default: 3.0 cm - ABNT standard)
            right: Right margin (default: 2.0 cm - ABNT standard)
            unit: Unit of measurement - 'cm' or 'inches' (default: 'cm')

        Raises:
            FileNotFoundError: If the Word file does not exist
            ValueError: If unit is invalid or any margin value is negative
            IOError: If the margins cannot be set

        Examples:
            # ABNT standard margins (Brazil academic standard)
            set_page_margins("tcc.docx", top=3.0, bottom=2.0, left=3.0, right=2.0)

            # Narrow margins for dense reports
            set_page_margins("relatorio.docx", top=1.5, bottom=1.5, left=1.5, right=1.5)

            # Inch-based margins
            set_page_margins("doc.docx", top=1.0, bottom=1.0, left=1.25, right=1.25, unit='inches')
        """
        logger.info(f"Setting page margins in {file_path}: top={top}, bottom={bottom}, left={left}, right={right} ({unit})")

        if unit not in ('cm', 'inches'):
            raise ValueError(f"Unidade inválida: '{unit}'. Use 'cm' ou 'inches'.")

        # Convert string values to float (Gemini API may send strings like "3" or "3cm")
        def _parse_margin(val):
            if isinstance(val, str):
                val = val.strip().lower()
                # Remove unit suffix if present (e.g., "3cm" -> "3", "2.5inches" -> "2.5")
                for suffix in ('cm', 'inches', 'inch', 'in'):
                    if val.endswith(suffix):
                        val = val[:-len(suffix)].strip()
                        break
                return float(val)
            return float(val)

        top = _parse_margin(top)
        bottom = _parse_margin(bottom)
        left = _parse_margin(left)
        right = _parse_margin(right)

        for val in (top, bottom, left, right):
            if val < 0:
                raise ValueError(f"Margens não podem ser negativas: {val}")

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = Document(file_path)

            converter = Cm if unit == 'cm' else Inches

            for section in doc.sections:
                section.top_margin = converter(top)
                section.bottom_margin = converter(bottom)
                section.left_margin = converter(left)
                section.right_margin = converter(right)

            doc.save(file_path)
            logger.info(f"Page margins set successfully ({unit})")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error setting page margins: {file_path} - {str(e)}")
            raise IOError(f"Failed to set page margins: {file_path} - {str(e)}") from e

    def set_page_size(self, file_path: str,
                      size: str = 'A4',
                      orientation: str = 'portrait') -> None:
        """Define o tamanho e orientação da página para todas as seções do documento.

        Tamanhos suportados: A4, A3, A5, Letter, Legal, Tabloid.
        Orientações: portrait (retrato) ou landscape (paisagem).

        Args:
            file_path: Path to the Word file
            size: Page size - 'A4', 'A3', 'A5', 'Letter', 'Legal', 'Tabloid' (default: 'A4')
            orientation: 'portrait' or 'landscape' (default: 'portrait')

        Raises:
            FileNotFoundError: If the Word file does not exist
            ValueError: If size or orientation is invalid
            IOError: If the page size cannot be set

        Examples:
            # Standard A4 portrait (most common)
            set_page_size("relatorio.docx", size="A4", orientation="portrait")

            # A4 landscape for wide tables
            set_page_size("planilha.docx", size="A4", orientation="landscape")

            # A3 for large diagrams
            set_page_size("diagrama.docx", size="A3", orientation="landscape")

            # US Letter for international documents
            set_page_size("doc.docx", size="Letter", orientation="portrait")
        """
        logger.info(f"Setting page size in {file_path}: size={size}, orientation={orientation}")

        # Page dimensions in cm (width x height in portrait)
        PAGE_SIZES = {
            'A4':      (21.0,  29.7),
            'A3':      (29.7,  42.0),
            'A5':      (14.8,  21.0),
            'Letter':  (21.59, 27.94),
            'Legal':   (21.59, 35.56),
            'Tabloid': (27.94, 43.18),
        }

        size_upper = size.upper() if size.upper() in {k.upper() for k in PAGE_SIZES} else size
        # Normalize case
        size_key = next((k for k in PAGE_SIZES if k.upper() == size.upper()), None)
        if size_key is None:
            raise ValueError(
                f"Tamanho de página inválido: '{size}'. "
                f"Opções: {', '.join(PAGE_SIZES.keys())}"
            )

        if orientation.lower() not in ('portrait', 'landscape'):
            raise ValueError(
                f"Orientação inválida: '{orientation}'. Use 'portrait' ou 'landscape'."
            )

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.enum.section import WD_ORIENT

            doc = Document(file_path)

            width_cm, height_cm = PAGE_SIZES[size_key]

            for section in doc.sections:
                if orientation.lower() == 'landscape':
                    section.orientation = WD_ORIENT.LANDSCAPE
                    # In landscape, swap width and height
                    section.page_width = Cm(height_cm)
                    section.page_height = Cm(width_cm)
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Cm(width_cm)
                    section.page_height = Cm(height_cm)

            doc.save(file_path)
            logger.info(f"Page size set to {size_key} {orientation} successfully")

        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error setting page size: {file_path} - {str(e)}")
            raise IOError(f"Failed to set page size: {file_path} - {str(e)}") from e

    def get_page_layout(self, file_path: str) -> Dict[str, Any]:
        """Retorna as configurações de layout de página do documento.

        Args:
            file_path: Path to the Word file

        Returns:
            Dictionary with page layout information:
            - page_width_cm, page_height_cm
            - orientation ('portrait' or 'landscape')
            - margins (top, bottom, left, right in cm)
            - section_count

        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Getting page layout from {file_path}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            doc = Document(file_path)
            section = doc.sections[0]

            width_cm = round(section.page_width.cm, 2)
            height_cm = round(section.page_height.cm, 2)
            orientation = 'landscape' if width_cm > height_cm else 'portrait'

            # Match page size name
            PAGE_SIZES = {
                'A4':      (21.0,  29.7),
                'A3':      (29.7,  42.0),
                'A5':      (14.8,  21.0),
                'Letter':  (21.59, 27.94),
                'Legal':   (21.59, 35.56),
                'Tabloid': (27.94, 43.18),
            }
            w_ref = min(width_cm, height_cm)
            h_ref = max(width_cm, height_cm)
            detected_size = 'Custom'
            for name, (pw, ph) in PAGE_SIZES.items():
                if abs(w_ref - pw) < 0.5 and abs(h_ref - ph) < 0.5:
                    detected_size = name
                    break

            result = {
                'page_size': detected_size,
                'page_width_cm': width_cm,
                'page_height_cm': height_cm,
                'orientation': orientation,
                'margins': {
                    'top_cm':    round(section.top_margin.cm, 2),
                    'bottom_cm': round(section.bottom_margin.cm, 2),
                    'left_cm':   round(section.left_margin.cm, 2),
                    'right_cm':  round(section.right_margin.cm, 2),
                },
                'section_count': len(doc.sections)
            }

            logger.info(f"Page layout: {detected_size} {orientation}, margins={result['margins']}")
            return result

        except Exception as e:
            logger.error(f"Error getting page layout: {str(e)}")
            raise

    # ==================== PAGE BREAK AND SECTION BREAK METHODS ====================

    def add_page_break(self, file_path: str, position: Optional[int] = None) -> None:
        """Insere uma quebra de página no documento.

        Por padrão adiciona a quebra ao final. Use 'position' para inserir
        após um parágrafo específico (índice 0-based).

        Args:
            file_path: Path to the Word file
            position: Paragraph index AFTER which to insert the break (0-based).
                      If None, appends to the end of the document.

        Raises:
            FileNotFoundError: If the Word file does not exist
            IndexError: If position is out of range
            IOError: If the page break cannot be added

        Examples:
            # Add page break at end of document
            add_page_break("relatorio.docx")

            # Add page break after paragraph 5
            add_page_break("doc.docx", position=5)
        """
        logger.info(f"Adding page break to {file_path}" + (f" after paragraph {position}" if position is not None else " at end"))

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.enum.text import WD_BREAK

            doc = Document(file_path)

            if position is not None:
                if position < 0 or position >= len(doc.paragraphs):
                    raise IndexError(
                        f"Índice de parágrafo inválido: {position}. "
                        f"Documento tem {len(doc.paragraphs)} parágrafos."
                    )
                # Insert a new paragraph with page break after the target paragraph
                target_para = doc.paragraphs[position]
                new_para = target_para.insert_paragraph_before()
                # Move it after target by re-inserting
                target_para._element.addnext(new_para._element)
                run = new_para.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                # Append page break at end
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_break(WD_BREAK.PAGE)

            doc.save(file_path)
            logger.info("Page break added successfully")

        except (IndexError, FileNotFoundError):
            raise
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding page break: {file_path} - {str(e)}")
            raise IOError(f"Failed to add page break: {file_path} - {str(e)}") from e

    def add_section_break(self, file_path: str,
                           break_type: str = "new_page",
                           position: Optional[int] = None) -> None:
        """Insere uma quebra de seção no documento.

        Quebras de seção permitem aplicar diferentes configurações de layout
        (margens, orientação, cabeçalho/rodapé) a partes distintas do documento.

        Args:
            file_path: Path to the Word file
            break_type: Type of section break:
                - 'new_page'   — próxima seção começa em nova página (padrão)
                - 'continuous' — seção contínua, sem quebra de página
                - 'even_page'  — próxima seção começa em página par
                - 'odd_page'   — próxima seção começa em página ímpar
            position: Paragraph index AFTER which to insert (0-based).
                      If None, appends to the end of the document.

        Raises:
            FileNotFoundError: If the Word file does not exist
            ValueError: If break_type is invalid
            IndexError: If position is out of range
            IOError: If the section break cannot be added

        Examples:
            # Add new-page section break at end
            add_section_break("relatorio.docx", break_type="new_page")

            # Add continuous section break after paragraph 3 (for column changes)
            add_section_break("doc.docx", break_type="continuous", position=3)

            # Add odd-page section break (for new chapters)
            add_section_break("livro.docx", break_type="odd_page")
        """
        logger.info(f"Adding '{break_type}' section break to {file_path}")

        BREAK_TYPES = {
            'new_page':   'nextPage',
            'continuous': 'continuous',
            'even_page':  'evenPage',
            'odd_page':   'oddPage',
        }

        if break_type.lower() not in BREAK_TYPES:
            raise ValueError(
                f"Tipo de quebra de seção inválido: '{break_type}'. "
                f"Opções: {', '.join(BREAK_TYPES.keys())}"
            )

        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")

        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document(file_path)

            if position is not None:
                if position < 0 or position >= len(doc.paragraphs):
                    raise IndexError(
                        f"Índice de parágrafo inválido: {position}. "
                        f"Documento tem {len(doc.paragraphs)} parágrafos."
                    )
                target_para = doc.paragraphs[position]
            else:
                # Add a new empty paragraph at the end to hold the sectPr
                target_para = doc.add_paragraph()

            # Add sectPr with the section break type to the paragraph's pPr
            p_elem = target_para._p
            ppr = p_elem.get_or_add_pPr()

            # Remove any existing sectPr in this paragraph
            existing = ppr.find(qn('w:sectPr'))
            if existing is not None:
                ppr.remove(existing)

            sect_pr = OxmlElement('w:sectPr')
            pg_sz = OxmlElement('w:pgSz')
            # Inherit page size from first section
            first_section = doc.sections[0]
            pg_sz.set(qn('w:w'), str(int(first_section.page_width.emu / 914400 * 1440)))
            pg_sz.set(qn('w:h'), str(int(first_section.page_height.emu / 914400 * 1440)))
            sect_pr.append(pg_sz)

            type_elem = OxmlElement('w:type')
            type_elem.set(qn('w:val'), BREAK_TYPES[break_type.lower()])
            sect_pr.append(type_elem)
            ppr.append(sect_pr)

            doc.save(file_path)
            logger.info(f"Section break '{break_type}' added successfully")

        except (IndexError, ValueError, FileNotFoundError):
            raise
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except Exception as e:
            logger.error(f"Error adding section break: {file_path} - {str(e)}")
            raise IOError(f"Failed to add section break: {file_path} - {str(e)}") from e

    def create_structured(self, file_path: str, elements: List[Dict[str, Any]]) -> None:
        """Cria documento Word estruturado com múltiplos tipos de elementos.
        
        Mais poderoso que create_word. Aceita uma lista de elementos:
        heading, paragraph, table, list.
        
        Args:
            file_path: Path where the Word file should be created
            elements: List of dicts, each with 'type' and type-specific fields:
                - {"type": "heading", "text": "...", "level": 1}
                - {"type": "paragraph", "text": "...", "bold": false, "italic": false}
                - {"type": "table", "headers": [...], "rows": [[...]]}
                - {"type": "list", "items": [...], "ordered": false}
            
        Raises:
            IOError: If the file cannot be created
        """
        logger.info(f"Creating structured Word file: {file_path} ({len(elements)} elements)")
        
        Path(file_path).parent.mkdir(parents=True, exist_ok=True)
        
        try:
            doc = Document()
            
            for elem in elements:
                elem_type = elem.get('type', 'paragraph')
                
                if elem_type == 'heading':
                    doc.add_heading(elem.get('text', ''), level=elem.get('level', 1))
                
                elif elem_type == 'paragraph':
                    p = doc.add_paragraph()
                    run = p.add_run(elem.get('text', ''))
                    if elem.get('bold'):
                        run.bold = True
                    if elem.get('italic'):
                        run.italic = True
                    if elem.get('font_size'):
                        run.font.size = Pt(elem['font_size'])
                    alignment = elem.get('alignment')
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    }
                    if alignment and alignment in align_map:
                        p.alignment = align_map[alignment]
                
                elif elem_type == 'table':
                    headers = elem.get('headers', [])
                    rows = elem.get('rows', [])
                    if headers:
                        total_rows = len(rows) + 1
                        total_cols = len(headers)
                        table = doc.add_table(rows=total_rows, cols=total_cols)
                        try:
                            table.style = elem.get('style', 'Table Grid')
                        except KeyError:
                            pass
                        for i, h in enumerate(headers):
                            cell = table.rows[0].cells[i]
                            cell.text = str(h)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        for r_idx, row_data in enumerate(rows):
                            for c_idx, val in enumerate(row_data):
                                if c_idx < total_cols:
                                    table.rows[r_idx + 1].cells[c_idx].text = str(val)
                
                elif elem_type == 'list':
                    items = elem.get('items', [])
                    ordered = elem.get('ordered', False)
                    style = 'List Number' if ordered else 'List Bullet'
                    for item in items:
                        doc.add_paragraph(item, style=style)
            
            doc.save(file_path)
            logger.info(f"Successfully created structured Word file with {len(elements)} elements")
            
        except Exception as e:
            logger.error(f"Error creating structured Word file: {file_path} - {str(e)}")
            raise IOError(f"Failed to create structured Word file: {file_path} - {str(e)}") from e

    # ==================== AI-POWERED TEXT IMPROVEMENT METHODS ====================
    
    def improve_text(self, file_path: str, improvement_type: str, 
                     target: str = "document", index: Optional[int] = None,
                     tone: Optional[str] = None) -> str:
        """Melhora texto usando IA do Gemini.
        
        Método genérico para melhorias de texto. Usado internamente por métodos específicos.
        
        Args:
            file_path: Path to the Word file
            improvement_type: Type of improvement ('grammar', 'clarity', 'tone', 'simplify', 'rewrite')
            target: 'document' for entire document or 'paragraph' for specific paragraph
            index: Paragraph index (required if target='paragraph')
            tone: Desired tone for 'tone' improvement type ('formal', 'informal', 'technical', 'casual')
            
        Returns:
            Improved text
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If gemini_client is not configured or parameters are invalid
            IndexError: If paragraph index is out of range
        """
        if not self.gemini_client:
            raise ValueError(
                "GeminiClient não configurado. Operações de melhoria de texto requerem "
                "integração com IA."
            )
        
        logger.info(f"Improving text in {file_path}: type={improvement_type}, target={target}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Get text to improve
            if target == "document":
                original_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if not original_text.strip():
                    raise ValueError("Documento está vazio ou contém apenas espaços em branco")
            elif target == "paragraph":
                if index is None:
                    raise ValueError("Index é obrigatório quando target='paragraph'")
                if index < 0 or index >= len(doc.paragraphs):
                    raise IndexError(
                        f"Paragraph index {index} out of range "
                        f"(document has {len(doc.paragraphs)} paragraphs)"
                    )
                original_text = doc.paragraphs[index].text
                if not original_text.strip():
                    raise ValueError(f"Parágrafo {index} está vazio")
            else:
                raise ValueError(f"Target inválido: {target}. Use 'document' ou 'paragraph'")
            
            # Build improvement prompt
            prompt = self._build_improvement_prompt(
                original_text, improvement_type, tone
            )
            
            # Get improved text from Gemini
            logger.info(f"Sending text to Gemini for {improvement_type} improvement")
            improved_text = self.gemini_client.generate_response(prompt, timeout=60)
            
            # Clean up response (remove markdown, extra formatting)
            improved_text = self._clean_ai_response(improved_text)
            
            # Validate improved text
            if not improved_text or len(improved_text.strip()) < 10:
                raise ValueError("IA retornou texto inválido ou muito curto")
            
            # Apply improved text
            if target == "document":
                # Replace all paragraphs
                # Clear existing paragraphs
                for para in doc.paragraphs:
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                
                # Add improved paragraphs
                for line in improved_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
            
            elif target == "paragraph":
                # Replace specific paragraph
                paragraph = doc.paragraphs[index]
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = improved_text
                else:
                    paragraph.add_run(improved_text)
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"Successfully improved text: {improvement_type}")
            return improved_text
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except (ValueError, IndexError):
            raise
        except Exception as e:
            logger.error(f"Error improving text: {file_path} - {str(e)}")
            raise IOError(f"Failed to improve text: {file_path} - {str(e)}") from e
    
    def _build_improvement_prompt(self, text: str, improvement_type: str, 
                                   tone: Optional[str] = None) -> str:
        """Constrói prompt especializado para melhoria de texto.
        
        Args:
            text: Original text to improve
            improvement_type: Type of improvement
            tone: Desired tone (for tone adjustments)
            
        Returns:
            Formatted prompt for Gemini
        """
        prompts = {
            'grammar': f"""Corrija APENAS os erros gramaticais no texto abaixo. 
Mantenha o estilo, tom e estrutura originais. Não adicione ou remova conteúdo.
Retorne APENAS o texto corrigido, sem explicações.

TEXTO ORIGINAL:
{text}

TEXTO CORRIGIDO:""",
            
            'clarity': f"""Melhore a clareza e coesão do texto abaixo.
Torne as ideias mais claras e fáceis de entender, mas mantenha o significado original.
Não mude o tom ou estilo. Retorne APENAS o texto melhorado, sem explicações.

TEXTO ORIGINAL:
{text}

TEXTO MELHORADO:""",
            
            'tone': f"""Ajuste o tom do texto abaixo para ser mais {tone or 'formal'}.
Mantenha o conteúdo e significado, mas adapte o estilo de escrita.
Retorne APENAS o texto ajustado, sem explicações.

TEXTO ORIGINAL:
{text}

TEXTO AJUSTADO:""",
            
            'simplify': f"""Simplifique a linguagem do texto abaixo para nível básico.
Use palavras simples e frases curtas. Mantenha todas as informações importantes.
Retorne APENAS o texto simplificado, sem explicações.

TEXTO ORIGINAL:
{text}

TEXTO SIMPLIFICADO:""",
            
            'rewrite': f"""Reescreva o texto abaixo de forma mais profissional e polida.
Melhore a estrutura, clareza e impacto, mas mantenha todas as informações.
Retorne APENAS o texto reescrito, sem explicações.

TEXTO ORIGINAL:
{text}

TEXTO REESCRITO:"""
        }
        
        return prompts.get(improvement_type, prompts['clarity'])
    
    def _clean_ai_response(self, response: str) -> str:
        """Limpa resposta da IA removendo markdown e formatação extra.
        
        Args:
            response: Raw response from AI
            
        Returns:
            Cleaned text
        """
        # Remove markdown code blocks
        response = re.sub(r'```[a-z]*\n?', '', response)
        response = re.sub(r'```', '', response)
        
        # Remove common AI prefixes
        prefixes = [
            'TEXTO CORRIGIDO:',
            'TEXTO MELHORADO:',
            'TEXTO AJUSTADO:',
            'TEXTO SIMPLIFICADO:',
            'TEXTO REESCRITO:',
            'Aqui está',
            'Segue',
        ]
        for prefix in prefixes:
            if response.strip().startswith(prefix):
                response = response.replace(prefix, '', 1)
        
        # Clean up extra whitespace
        response = response.strip()
        
        return response
    
    # Convenience methods for specific improvements
    
    def correct_grammar(self, file_path: str, target: str = "document", 
                       index: Optional[int] = None) -> str:
        """Corrige gramática do documento ou parágrafo específico.
        
        Args:
            file_path: Path to the Word file
            target: 'document' or 'paragraph'
            index: Paragraph index (if target='paragraph')
            
        Returns:
            Corrected text
        """
        return self.improve_text(file_path, 'grammar', target, index)
    
    def improve_clarity(self, file_path: str, target: str = "document",
                       index: Optional[int] = None) -> str:
        """Melhora clareza e coesão do texto.
        
        Args:
            file_path: Path to the Word file
            target: 'document' or 'paragraph'
            index: Paragraph index (if target='paragraph')
            
        Returns:
            Improved text
        """
        return self.improve_text(file_path, 'clarity', target, index)
    
    def adjust_tone(self, file_path: str, tone: str, target: str = "document",
                   index: Optional[int] = None) -> str:
        """Ajusta o tom do texto.
        
        Args:
            file_path: Path to the Word file
            tone: Desired tone ('formal', 'informal', 'technical', 'casual')
            target: 'document' or 'paragraph'
            index: Paragraph index (if target='paragraph')
            
        Returns:
            Text with adjusted tone
        """
        return self.improve_text(file_path, 'tone', target, index, tone=tone)
    
    def simplify_language(self, file_path: str, target: str = "document",
                         index: Optional[int] = None) -> str:
        """Simplifica a linguagem para nível básico.
        
        Args:
            file_path: Path to the Word file
            target: 'document' or 'paragraph'
            index: Paragraph index (if target='paragraph')
            
        Returns:
            Simplified text
        """
        return self.improve_text(file_path, 'simplify', target, index)
    
    def rewrite_professional(self, file_path: str, target: str = "document",
                            index: Optional[int] = None) -> str:
        """Reescreve o texto de forma mais profissional.
        
        Args:
            file_path: Path to the Word file
            target: 'document' or 'paragraph'
            index: Paragraph index (if target='paragraph')
            
        Returns:
            Professionally rewritten text
        """
        return self.improve_text(file_path, 'rewrite', target, index)

    # ==================== AI-POWERED CONTENT ANALYSIS AND GENERATION ====================
    
    def analyze_and_generate(self, file_path: str, analysis_type: str,
                            output_mode: str = "new_section", 
                            section_title: Optional[str] = None,
                            size: Optional[str] = None,
                            num_items: int = 5) -> str:
        """Analisa conteúdo e gera novo conteúdo baseado na análise.
        
        Método genérico para análise e geração de conteúdo. Usado internamente por métodos específicos.
        
        Args:
            file_path: Path to the Word file
            analysis_type: Type of analysis ('summary', 'key_points', 'resume', 'conclusions', 'faq')
            output_mode: Where to add generated content:
                - 'new_section': Add as new section with heading
                - 'append': Append to end of document
                - 'new_document': Create new document (not implemented yet)
            section_title: Title for new section (if output_mode='new_section')
            size: Size for resume ('1_page', '1_paragraph', '3_sentences')
            num_items: Number of items to generate (for key_points, conclusions, faq)
            
        Returns:
            Generated content
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If gemini_client is not configured or parameters are invalid
        """
        if not self.gemini_client:
            raise ValueError(
                "GeminiClient não configurado. Operações de análise requerem "
                "integração com IA."
            )
        
        # Validate output_mode
        valid_modes = ['new_section', 'append']
        if output_mode not in valid_modes:
            raise ValueError(f"Output mode inválido: '{output_mode}'. Use: {', '.join(valid_modes)}")
        
        logger.info(f"Analyzing and generating content: type={analysis_type}, mode={output_mode}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            # Read document content
            doc = Document(file_path)
            original_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not original_text.strip():
                raise ValueError("Documento está vazio ou contém apenas espaços em branco")
            
            # Build analysis prompt
            prompt = self._build_analysis_prompt(
                original_text, analysis_type, size, num_items
            )
            
            # Get generated content from Gemini
            logger.info(f"Sending text to Gemini for {analysis_type} analysis")
            generated_content = self.gemini_client.generate_response(prompt, timeout=60)
            
            # Clean up response
            generated_content = self._clean_ai_response(generated_content)
            
            # Validate generated content
            if not generated_content or len(generated_content.strip()) < 20:
                raise ValueError("IA retornou conteúdo inválido ou muito curto")
            
            # Add generated content to document
            if output_mode == "new_section":
                # Add heading
                title = section_title or self._get_default_section_title(analysis_type)
                doc.add_heading(title, level=1)
                
                # Add content
                if analysis_type in ['key_points', 'faq']:
                    # Add as list
                    self._add_formatted_content(doc, generated_content, analysis_type)
                else:
                    # Add as paragraphs
                    for line in generated_content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
            
            elif output_mode == "append":
                # Just append to end
                if analysis_type in ['key_points', 'faq']:
                    self._add_formatted_content(doc, generated_content, analysis_type)
                else:
                    for line in generated_content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
            
            else:
                raise ValueError(f"Output mode inválido: {output_mode}")
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"Successfully generated {analysis_type} content")
            return generated_content
            
        except PackageNotFoundError as e:
            raise CorruptedFileError(f"Word file is corrupted or invalid: {file_path}") from e
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error analyzing and generating content: {file_path} - {str(e)}")
            raise IOError(f"Failed to analyze and generate content: {file_path} - {str(e)}") from e
    
    def _build_analysis_prompt(self, text: str, analysis_type: str,
                               size: Optional[str] = None, num_items: int = 5) -> str:
        """Constrói prompt especializado para análise e geração de conteúdo.
        
        Args:
            text: Original text to analyze
            analysis_type: Type of analysis
            size: Size for resume
            num_items: Number of items to generate
            
        Returns:
            Formatted prompt for Gemini
        """
        prompts = {
            'summary': f"""Analise o texto abaixo e crie um SUMÁRIO EXECUTIVO de no máximo 1 página.
O sumário deve:
- Capturar os pontos mais importantes
- Ser claro e objetivo
- Usar linguagem profissional
- Ter entre 200-300 palavras

Retorne APENAS o sumário, sem títulos ou explicações.

TEXTO ORIGINAL:
{text}

SUMÁRIO EXECUTIVO:""",
            
            'key_points': f"""Analise o texto abaixo e extraia os {num_items} PONTOS-CHAVE mais importantes.

Retorne APENAS uma lista numerada dos pontos, sem introdução ou conclusão.
Cada ponto deve ser claro, conciso e capturar uma ideia principal.

TEXTO ORIGINAL:
{text}

PONTOS-CHAVE:""",
            
            'resume': self._build_resume_prompt(text, size),
            
            'conclusions': f"""Analise o texto abaixo e identifique as {num_items} CONCLUSÕES PRINCIPAIS.

As conclusões devem:
- Representar insights ou aprendizados principais
- Ser baseadas no conteúdo do texto
- Ser claras e objetivas

Retorne APENAS uma lista numerada das conclusões, sem introdução.

TEXTO ORIGINAL:
{text}

CONCLUSÕES PRINCIPAIS:""",
            
            'faq': f"""Analise o texto abaixo e crie uma seção de FAQ (Perguntas Frequentes) com {num_items} perguntas e respostas.

As perguntas devem:
- Ser relevantes ao conteúdo
- Cobrir os pontos principais
- Ser perguntas que um leitor realmente faria

Formato:
P: [pergunta]
R: [resposta]

Retorne APENAS as perguntas e respostas, sem introdução.

TEXTO ORIGINAL:
{text}

FAQ:"""
        }
        
        return prompts.get(analysis_type, prompts['summary'])
    
    def _build_resume_prompt(self, text: str, size: Optional[str]) -> str:
        """Constrói prompt específico para resumos de diferentes tamanhos."""
        size_instructions = {
            '1_page': 'no máximo 1 página (200-300 palavras)',
            '1_paragraph': 'em exatamente 1 parágrafo (50-100 palavras)',
            '3_sentences': 'em exatamente 3 frases curtas e diretas'
        }
        
        instruction = size_instructions.get(size, size_instructions['1_paragraph'])
        
        return f"""Crie um RESUMO do texto abaixo {instruction}.

O resumo deve:
- Capturar as informações mais importantes
- Ser claro e objetivo
- Manter o significado original

Retorne APENAS o resumo, sem título ou explicações.

TEXTO ORIGINAL:
{text}

RESUMO:"""
    
    def _get_default_section_title(self, analysis_type: str) -> str:
        """Retorna título padrão para seção baseado no tipo de análise."""
        titles = {
            'summary': 'Sumário Executivo',
            'key_points': 'Pontos-Chave',
            'resume': 'Resumo',
            'conclusions': 'Conclusões',
            'faq': 'Perguntas Frequentes (FAQ)'
        }
        return titles.get(analysis_type, 'Análise')
    
    def _add_formatted_content(self, doc: Document, content: str, content_type: str) -> None:
        """Adiciona conteúdo formatado ao documento (listas, FAQ, etc)."""
        if content_type == 'key_points':
            # Parse numbered list
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    # Remove numbering/bullets
                    clean_line = re.sub(r'^\d+[\.\)]\s*', '', line)
                    clean_line = re.sub(r'^[-•]\s*', '', clean_line)
                    if clean_line:
                        doc.add_paragraph(clean_line, style='List Number')
        
        elif content_type == 'faq':
            # Parse FAQ format (P: ... R: ...)
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('P:') or line.startswith('Q:'):
                    # Question - bold
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                elif line.startswith('R:') or line.startswith('A:'):
                    # Answer - normal
                    doc.add_paragraph(line)
                elif line:
                    # Other content
                    doc.add_paragraph(line)
    
    # Convenience methods for specific analyses
    
    def generate_summary(self, file_path: str, output_mode: str = "new_section",
                        section_title: Optional[str] = None) -> str:
        """Gera sumário executivo do documento.
        
        Args:
            file_path: Path to the Word file
            output_mode: 'new_section' or 'append'
            section_title: Custom title for section
            
        Returns:
            Generated summary
        """
        return self.analyze_and_generate(
            file_path, 'summary', output_mode, section_title
        )
    
    def extract_key_points(self, file_path: str, num_points: int = 5,
                          output_mode: str = "new_section",
                          section_title: Optional[str] = None) -> str:
        """Extrai pontos-chave do documento.
        
        Args:
            file_path: Path to the Word file
            num_points: Number of key points to extract
            output_mode: 'new_section' or 'append'
            section_title: Custom title for section
            
        Returns:
            Generated key points
        """
        return self.analyze_and_generate(
            file_path, 'key_points', output_mode, section_title, num_items=num_points
        )
    
    def create_resume(self, file_path: str, size: str = "1_paragraph",
                     output_mode: str = "new_section",
                     section_title: Optional[str] = None) -> str:
        """Cria resumo do documento em tamanho especificado.
        
        Args:
            file_path: Path to the Word file
            size: '1_page', '1_paragraph', or '3_sentences'
            output_mode: 'new_section' or 'append'
            section_title: Custom title for section
            
        Returns:
            Generated resume
        """
        return self.analyze_and_generate(
            file_path, 'resume', output_mode, section_title, size=size
        )
    
    def generate_conclusions(self, file_path: str, num_conclusions: int = 3,
                            output_mode: str = "new_section",
                            section_title: Optional[str] = None) -> str:
        """Gera conclusões principais do documento.
        
        Args:
            file_path: Path to the Word file
            num_conclusions: Number of conclusions to generate
            output_mode: 'new_section' or 'append'
            section_title: Custom title for section
            
        Returns:
            Generated conclusions
        """
        return self.analyze_and_generate(
            file_path, 'conclusions', output_mode, section_title, num_items=num_conclusions
        )
    
    def create_faq(self, file_path: str, num_questions: int = 5,
                  output_mode: str = "new_section",
                  section_title: Optional[str] = None) -> str:
        """Cria seção de FAQ baseada no conteúdo do documento.
        
        Args:
            file_path: Path to the Word file
            num_questions: Number of Q&A pairs to generate
            output_mode: 'new_section' or 'append'
            section_title: Custom title for section
            
        Returns:
            Generated FAQ
        """
        return self.analyze_and_generate(
            file_path, 'faq', output_mode, section_title, num_items=num_questions
        )

    # ==================== CONVERSION METHODS ====================
    # Phase 3: Format Conversion and Transformation
    
    def convert_list_to_table(self, file_path: str, list_index: int = 0,
                              include_header: bool = False,
                              header_text: Optional[str] = None) -> str:
        """Converte lista numerada ou bullet em tabela.
        
        Args:
            file_path: Path to the Word file
            list_index: Index of the list to convert (0 = first list)
            include_header: Whether to add a header row
            header_text: Text for header (default: "Item")
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If list_index is invalid or no lists found
        """
        logger.info(f"Converting list {list_index} to table in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Find all list items
            list_groups = self._find_list_groups(doc)
            
            if not list_groups:
                raise ValueError("Nenhuma lista encontrada no documento")
            
            if list_index < 0 or list_index >= len(list_groups):
                raise ValueError(
                    f"Índice de lista inválido: {list_index}. "
                    f"Documento tem {len(list_groups)} lista(s)"
                )
            
            # Get the target list
            list_items = list_groups[list_index]
            
            # Find position to insert table (before first list item)
            first_para = list_items[0]
            insert_index = None
            for i, para in enumerate(doc.paragraphs):
                if para == first_para:
                    insert_index = i
                    break
            
            # Create table
            num_rows = len(list_items) + (1 if include_header else 0)
            table = doc.add_table(rows=num_rows, cols=1)
            table.style = 'Light Grid Accent 1'
            
            # Add header if requested
            row_offset = 0
            if include_header:
                header_cell = table.rows[0].cells[0]
                header_cell.text = header_text or "Item"
                header_cell.paragraphs[0].runs[0].bold = True
                row_offset = 1
            
            # Fill table with list items
            for i, para in enumerate(list_items):
                # Remove list formatting markers (numbers, bullets)
                text = para.text.strip()
                # Remove leading numbers/bullets
                import re
                text = re.sub(r'^\d+[\.\)]\s*', '', text)  # Remove "1. " or "1) "
                text = re.sub(r'^[•\-\*]\s*', '', text)     # Remove "• " or "- " or "* "
                
                table.rows[i + row_offset].cells[0].text = text
            
            # Remove original list paragraphs
            for para in reversed(list_items):
                p_element = para._element
                p_element.getparent().remove(p_element)
            
            # Move table to correct position
            if insert_index is not None:
                table_element = table._element
                doc._element.body.insert(insert_index, table_element)
            
            doc.save(file_path)
            
            message = f"Lista {list_index} convertida em tabela com {len(list_items)} itens"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error converting list to table: {str(e)}")
            raise
    
    def convert_table_to_list(self, file_path: str, table_index: int = 0,
                             list_type: str = "numbered",
                             skip_header: bool = False,
                             separator: str = " - ") -> str:
        """Converte tabela em lista numerada ou bullet.
        
        Args:
            file_path: Path to the Word file
            table_index: Index of the table to convert (0 = first table)
            list_type: 'numbered' or 'bullet'
            skip_header: Skip first row (if it's a header)
            separator: Separator for multiple columns (default: " - ")
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If table_index is invalid or no tables found
        """
        logger.info(f"Converting table {table_index} to {list_type} list in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        if list_type not in ['numbered', 'bullet']:
            raise ValueError(f"Tipo de lista inválido: '{list_type}'. Use 'numbered' ou 'bullet'")
        
        try:
            doc = Document(file_path)
            
            if not doc.tables:
                raise ValueError("Nenhuma tabela encontrada no documento")
            
            if table_index < 0 or table_index >= len(doc.tables):
                raise ValueError(
                    f"Índice de tabela inválido: {table_index}. "
                    f"Documento tem {len(doc.tables)} tabela(s)"
                )
            
            # Get the target table
            table = doc.tables[table_index]
            
            # Find position to insert list (before table)
            table_element = table._element
            insert_index = None
            for i, element in enumerate(doc._element.body):
                if element == table_element:
                    insert_index = i
                    break
            
            # Extract table data
            start_row = 1 if skip_header else 0
            list_items = []
            
            for row in table.rows[start_row:]:
                # Combine all cells in row
                cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    item_text = separator.join(cell_texts)
                    list_items.append(item_text)
            
            # Remove table
            table_element.getparent().remove(table_element)
            
            # Add list items
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.shared import Pt
            
            for i, item_text in enumerate(list_items):
                para = doc.add_paragraph()
                
                if list_type == 'numbered':
                    para.style = 'List Number'
                    para.text = item_text
                else:  # bullet
                    para.style = 'List Bullet'
                    para.text = item_text
                
                # Move paragraph to correct position
                if insert_index is not None:
                    para_element = para._element
                    doc._element.body.insert(insert_index + i, para_element)
            
            doc.save(file_path)
            
            message = f"Tabela {table_index} convertida em lista {list_type} com {len(list_items)} itens"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error converting table to list: {str(e)}")
            raise
    
    def _find_list_groups(self, doc: Document) -> List[List]:
        """Encontra grupos de parágrafos que formam listas.
        
        Args:
            doc: Document object
            
        Returns:
            List of list groups, where each group is a list of paragraphs
        """
        import re
        
        list_groups = []
        current_group = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check if paragraph is a list item
            is_numbered = bool(re.match(r'^\d+[\.\)]\s+', text))
            is_bullet = bool(re.match(r'^[•\-\*]\s+', text))
            is_list_style = para.style.name in ['List Number', 'List Bullet', 'List Paragraph']
            
            if is_numbered or is_bullet or is_list_style:
                current_group.append(para)
            else:
                # End of current list group
                if current_group:
                    list_groups.append(current_group)
                    current_group = []
        
        # Add last group if exists
        if current_group:
            list_groups.append(current_group)
        
        return list_groups

    def extract_tables_to_excel(self, file_path: str, output_path: str,
                                sheet_names: Optional[List[str]] = None) -> str:
        """Extrai todas as tabelas do Word para um arquivo Excel.
        
        Cada tabela é colocada em uma aba separada do Excel.
        
        Args:
            file_path: Path to the Word file
            output_path: Path for the output Excel file
            sheet_names: Optional list of names for Excel sheets
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the Word file does not exist
            ValueError: If no tables found or sheet_names length mismatch
        """
        logger.info(f"Extracting tables from {file_path} to Excel: {output_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            # Extract tables from Word
            tables_data = self.extract_tables(file_path)
            
            if not tables_data:
                raise ValueError("Nenhuma tabela encontrada no documento Word")
            
            # Validate sheet_names if provided
            if sheet_names:
                if len(sheet_names) != len(tables_data):
                    raise ValueError(
                        f"Número de nomes de abas ({len(sheet_names)}) não corresponde "
                        f"ao número de tabelas ({len(tables_data)})"
                    )
            else:
                # Generate default sheet names
                sheet_names = [f"Tabela {i+1}" for i in range(len(tables_data))]
            
            # Create Excel workbook with tables
            from openpyxl import Workbook
            
            wb = Workbook()
            # Remove default sheet
            if 'Sheet' in wb.sheetnames:
                wb.remove(wb['Sheet'])
            
            for i, (table_data, sheet_name) in enumerate(zip(tables_data, sheet_names)):
                ws = wb.create_sheet(title=sheet_name)
                
                # Write table data to worksheet
                for row_idx, row_data in enumerate(table_data, start=1):
                    for col_idx, cell_value in enumerate(row_data, start=1):
                        ws.cell(row=row_idx, column=col_idx, value=cell_value)
                
                # Auto-adjust column widths
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save Excel file
            wb.save(output_path)
            
            message = f"Extraídas {len(tables_data)} tabela(s) para {output_path}"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error extracting tables to Excel: {str(e)}")
            raise

    def export_to_txt(self, file_path: str, output_path: str) -> str:
        """Exporta documento Word para texto puro (.txt).
        
        Args:
            file_path: Path to the Word file
            output_path: Path for the output TXT file
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Exporting {file_path} to TXT: {output_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract all text
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Write to TXT file
            output = Path(output_path)
            output.write_text('\n\n'.join(text_content), encoding='utf-8')
            
            message = f"Documento exportado para TXT: {output_path}"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error exporting to TXT: {str(e)}")
            raise
    
    def export_to_markdown(self, file_path: str, output_path: str) -> str:
        """Exporta documento Word para Markdown (.md).
        
        Converte:
        - Títulos (Heading 1-6) para # ## ### etc
        - Listas numeradas e bullet points
        - Tabelas para formato Markdown
        - Negrito e itálico
        
        Args:
            file_path: Path to the Word file
            output_path: Path for the output Markdown file
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Exporting {file_path} to Markdown: {output_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            markdown_lines = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    markdown_lines.append('')
                    continue
                
                # Convert headings
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        markdown_lines.append(f"{'#' * level_num} {para.text}")
                    except ValueError:
                        markdown_lines.append(para.text)
                
                # Convert lists
                elif para.style.name == 'List Number':
                    markdown_lines.append(f"1. {para.text}")
                elif para.style.name in ['List Bullet', 'List Paragraph']:
                    markdown_lines.append(f"- {para.text}")
                
                # Regular paragraph
                else:
                    # Convert formatting
                    text = self._convert_runs_to_markdown(para)
                    markdown_lines.append(text)
            
            # Process tables
            for table in doc.tables:
                markdown_lines.append('')
                markdown_lines.extend(self._table_to_markdown(table))
                markdown_lines.append('')
            
            # Write to Markdown file
            output = Path(output_path)
            output.write_text('\n'.join(markdown_lines), encoding='utf-8')
            
            message = f"Documento exportado para Markdown: {output_path}"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error exporting to Markdown: {str(e)}")
            raise
    
    def _convert_runs_to_markdown(self, para) -> str:
        """Converte runs de parágrafo para Markdown com formatação.
        
        Args:
            para: Paragraph object
            
        Returns:
            Formatted Markdown text
        """
        result = []
        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result.append(text)
        return ''.join(result)
    
    def _table_to_markdown(self, table) -> List[str]:
        """Converte tabela Word para formato Markdown.
        
        Args:
            table: Table object
            
        Returns:
            List of Markdown table lines
        """
        lines = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
            
            # Add separator after first row (header)
            if i == 0:
                separators = ['---'] * len(cells)
                lines.append('| ' + ' | '.join(separators) + ' |')
        
        return lines
    
    def export_to_html(self, file_path: str, output_path: str) -> str:
        """Exporta documento Word para HTML (.html).
        
        Converte:
        - Títulos para <h1> <h2> etc
        - Parágrafos para <p>
        - Listas para <ul> <ol>
        - Tabelas para <table>
        - Negrito e itálico para <b> <i>
        
        Args:
            file_path: Path to the Word file
            output_path: Path for the output HTML file
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the Word file does not exist
        """
        logger.info(f"Exporting {file_path} to HTML: {output_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            html_lines = [
                '<!DOCTYPE html>',
                '<html>',
                '<head>',
                '    <meta charset="UTF-8">',
                '    <title>Document</title>',
                '    <style>',
                '        body { font-family: Arial, sans-serif; margin: 40px; }',
                '        table { border-collapse: collapse; width: 100%; margin: 20px 0; }',
                '        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }',
                '        th { background-color: #f2f2f2; }',
                '    </style>',
                '</head>',
                '<body>'
            ]
            
            # Process paragraphs with proper list grouping
            in_ordered_list = False
            in_unordered_list = False

            def _close_open_list():
                if in_ordered_list:
                    html_lines.append('    </ol>')
                elif in_unordered_list:
                    html_lines.append('    </ul>')

            for para in doc.paragraphs:
                if not para.text.strip():
                    _close_open_list()
                    in_ordered_list = False
                    in_unordered_list = False
                    continue
                
                # Convert headings
                if para.style.name.startswith('Heading'):
                    _close_open_list()
                    in_ordered_list = False
                    in_unordered_list = False
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        html_lines.append(f"    <h{level_num}>{para.text}</h{level_num}>")
                    except ValueError:
                        html_lines.append(f"    <p>{para.text}</p>")
                
                # Convert numbered list items
                elif para.style.name == 'List Number':
                    if in_unordered_list:
                        html_lines.append('    </ul>')
                        in_unordered_list = False
                    if not in_ordered_list:
                        html_lines.append('    <ol>')
                        in_ordered_list = True
                    html_lines.append(f"        <li>{para.text}</li>")
                
                # Convert bullet list items
                elif para.style.name in ['List Bullet', 'List Paragraph']:
                    if in_ordered_list:
                        html_lines.append('    </ol>')
                        in_ordered_list = False
                    if not in_unordered_list:
                        html_lines.append('    <ul>')
                        in_unordered_list = True
                    html_lines.append(f"        <li>{para.text}</li>")
                
                # Regular paragraph
                else:
                    _close_open_list()
                    in_ordered_list = False
                    in_unordered_list = False
                    text = self._convert_runs_to_html(para)
                    html_lines.append(f"    <p>{text}</p>")

            # Close any still-open list at end of paragraphs
            _close_open_list()

            # Process tables
            for table in doc.tables:
                html_lines.extend(self._table_to_html(table))
            
            html_lines.extend([
                '</body>',
                '</html>'
            ])
            
            # Write to HTML file
            output = Path(output_path)
            output.write_text('\n'.join(html_lines), encoding='utf-8')
            
            message = f"Documento exportado para HTML: {output_path}"
            logger.info(message)
            return message
            
        except Exception as e:
            logger.error(f"Error exporting to HTML: {str(e)}")
            raise
    
    def _convert_runs_to_html(self, para) -> str:
        """Converte runs de parágrafo para HTML com formatação.
        
        Args:
            para: Paragraph object
            
        Returns:
            Formatted HTML text
        """
        result = []
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            result.append(text)
        return ''.join(result)
    
    def _table_to_html(self, table) -> List[str]:
        """Converte tabela Word para HTML.
        
        Args:
            table: Table object
            
        Returns:
            List of HTML table lines
        """
        lines = ['    <table>']
        
        for i, row in enumerate(table.rows):
            lines.append('        <tr>')
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                lines.append(f'            <{tag}>{cell.text}</{tag}>')
            lines.append('        </tr>')
        
        lines.append('    </table>')
        return lines
    
    def export_to_pdf(self, file_path: str, output_path: str) -> str:
        """Exporta documento Word para PDF (.pdf).
        
        Requer docx2pdf instalado (Windows) ou LibreOffice (Linux/Mac).
        
        Args:
            file_path: Path to the Word file
            output_path: Path for the output PDF file
            
        Returns:
            Success message
            
        Raises:
            FileNotFoundError: If the Word file does not exist
            ImportError: If docx2pdf is not installed
            RuntimeError: If conversion fails
        """
        logger.info(f"Exporting {file_path} to PDF: {output_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            # Try to import docx2pdf
            try:
                from docx2pdf import convert
            except ImportError:
                raise ImportError(
                    "docx2pdf não está instalado. "
                    "Instale com: pip install docx2pdf"
                )
            
            # Convert to PDF
            convert(str(file_path), str(output_path))
            
            message = f"Documento exportado para PDF: {output_path}"
            logger.info(message)
            return message
            
        except ImportError:
            raise
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            raise RuntimeError(f"Falha ao converter para PDF: {str(e)}") from e

    # ==================== DOCUMENT ANALYSIS METHODS ====================
    # Phase 4: Document Analysis and Insights
    
    def analyze_word_count(self, file_path: str) -> Dict[str, Any]:
        """Analisa contagem de palavras total e por seção.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with word count statistics
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Analyzing word count in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            total_words = 0
            total_paragraphs = 0
            sections = []
            current_section = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if it's a heading (new section)
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_section:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        "title": text,
                        "level": para.style.name,
                        "words": 0,
                        "paragraphs": 0
                    }
                else:
                    # Count words in paragraph
                    words = len(text.split())
                    total_words += words
                    total_paragraphs += 1
                    
                    if current_section:
                        current_section["words"] += words
                        current_section["paragraphs"] += 1
            
            # Save last section
            if current_section:
                sections.append(current_section)
            
            # Calculate percentages
            for section in sections:
                section["percentage"] = round((section["words"] / total_words * 100), 2) if total_words > 0 else 0
            
            result = {
                "total_words": total_words,
                "total_paragraphs": total_paragraphs,
                "average_words_per_paragraph": round(total_words / total_paragraphs, 2) if total_paragraphs > 0 else 0,
                "sections": sections,
                "section_count": len(sections)
            }
            
            logger.info(f"Word count analysis complete: {total_words} words, {len(sections)} sections")
            return result
            
        except Exception as e:
            logger.error(f"Error analyzing word count: {str(e)}")
            raise
    
    def analyze_section_length(self, file_path: str, max_words: int = 500) -> Dict[str, Any]:
        """Identifica seções muito longas no documento.
        
        Args:
            file_path: Path to the Word file
            max_words: Maximum words per section (default: 500)
            
        Returns:
            Dictionary with long sections analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Analyzing section length in: {file_path}")
        
        # Get word count analysis
        word_count = self.analyze_word_count(file_path)
        
        long_sections = []
        for section in word_count["sections"]:
            if section["words"] > max_words:
                long_sections.append({
                    "title": section["title"],
                    "words": section["words"],
                    "excess_words": section["words"] - max_words,
                    "recommendation": f"Considere dividir em sub-seções ou reduzir em {section['words'] - max_words} palavras"
                })
        
        result = {
            "max_words_threshold": max_words,
            "long_sections_count": len(long_sections),
            "long_sections": long_sections,
            "total_sections": word_count["section_count"],
            "percentage_long": round((len(long_sections) / word_count["section_count"] * 100), 2) if word_count["section_count"] > 0 else 0
        }
        
        logger.info(f"Section length analysis complete: {len(long_sections)} long sections found")
        return result
    
    def get_document_statistics(self, file_path: str) -> Dict[str, Any]:
        """Obtém estatísticas gerais do documento.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with document statistics
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Getting document statistics for: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Count elements
            total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            total_tables = len(doc.tables)
            total_headings = len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])
            
            # Count words and characters
            all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            total_words = len(all_text.split())
            total_characters = len(all_text)
            total_characters_no_spaces = len(all_text.replace(' ', ''))
            
            # Count sentences (approximate)
            import re
            sentences = re.split(r'[.!?]+', all_text)
            total_sentences = len([s for s in sentences if s.strip()])
            
            # Calculate averages
            avg_words_per_sentence = round(total_words / total_sentences, 2) if total_sentences > 0 else 0
            avg_chars_per_word = round(total_characters_no_spaces / total_words, 2) if total_words > 0 else 0
            
            result = {
                "total_paragraphs": total_paragraphs,
                "total_words": total_words,
                "total_characters": total_characters,
                "total_characters_no_spaces": total_characters_no_spaces,
                "total_sentences": total_sentences,
                "total_tables": total_tables,
                "total_headings": total_headings,
                "average_words_per_sentence": avg_words_per_sentence,
                "average_characters_per_word": avg_chars_per_word,
                "estimated_reading_time_minutes": round(total_words / 200, 1)  # 200 words per minute
            }
            
            logger.info(f"Document statistics complete: {total_words} words, {total_paragraphs} paragraphs")
            return result
            
        except Exception as e:
            logger.error(f"Error getting document statistics: {str(e)}")
            raise
    
    def analyze_tone(self, file_path: str) -> Dict[str, Any]:
        """Analisa o tom do documento usando IA.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with tone analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If gemini_client is not configured
        """
        if not self.gemini_client:
            raise ValueError(
                "GeminiClient não configurado. Análise de tom requer integração com IA."
            )
        
        logger.info(f"Analyzing tone in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text.strip():
                raise ValueError("Documento está vazio")
            
            # Build prompt for tone analysis
            prompt = f"""Analise o tom do seguinte texto e classifique-o.

Texto:
{text[:3000]}

Forneça uma análise estruturada com:
1. Tom principal (formal, informal, técnico, casual)
2. Pontuação de 0 a 1 para cada categoria de tom
3. Análise textual explicando o tom identificado
4. Recomendações para manter ou ajustar o tom

Responda em formato claro e estruturado."""
            
            logger.info("Sending text to Gemini for tone analysis")
            response = self.gemini_client.generate_response(prompt, timeout=60)
            
            # Parse response (simplified - in production would use structured output)
            result = {
                "analysis": response,
                "text_sample": text[:500] + "..." if len(text) > 500 else text
            }
            
            logger.info("Tone analysis complete")
            return result
            
        except Exception as e:
            logger.error(f"Error analyzing tone: {str(e)}")
            raise
    
    def identify_jargon(self, file_path: str) -> Dict[str, Any]:
        """Identifica jargões técnicos e termos complexos usando IA.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with jargon analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If gemini_client is not configured
        """
        if not self.gemini_client:
            raise ValueError(
                "GeminiClient não configurado. Identificação de jargões requer integração com IA."
            )
        
        logger.info(f"Identifying jargon in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text.strip():
                raise ValueError("Documento está vazio")
            
            # Build prompt for jargon identification
            prompt = f"""Identifique jargões técnicos, termos complexos ou palavras difíceis no seguinte texto.

Texto:
{text[:3000]}

Para cada termo identificado, forneça:
1. O termo/jargão
2. Contexto onde aparece
3. Alternativas mais simples
4. Explicação de por que é considerado jargão

Liste os termos em ordem de complexidade (mais complexo primeiro).
Limite a 10 termos principais."""
            
            logger.info("Sending text to Gemini for jargon identification")
            response = self.gemini_client.generate_response(prompt, timeout=60)
            
            # Parse response
            result = {
                "analysis": response,
                "text_sample": text[:500] + "..." if len(text) > 500 else text
            }
            
            logger.info("Jargon identification complete")
            return result
            
        except Exception as e:
            logger.error(f"Error identifying jargon: {str(e)}")
            raise
    
    def analyze_readability(self, file_path: str) -> Dict[str, Any]:
        """Analisa a legibilidade do documento.
        
        Calcula métricas de legibilidade e fornece análise com IA.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with readability analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Analyzing readability in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text.strip():
                raise ValueError("Documento está vazio")
            
            # Get basic statistics
            stats = self.get_document_statistics(file_path)
            
            # Calculate readability metrics
            words = stats["total_words"]
            sentences = stats["total_sentences"]
            
            # Simplified readability score (0-100, higher = easier)
            # Based on average sentence length
            avg_sentence_length = stats["average_words_per_sentence"]
            
            if avg_sentence_length < 10:
                readability_score = 90
                reading_level = "Muito Fácil"
            elif avg_sentence_length < 15:
                readability_score = 75
                reading_level = "Fácil"
            elif avg_sentence_length < 20:
                readability_score = 60
                reading_level = "Moderado"
            elif avg_sentence_length < 25:
                readability_score = 45
                reading_level = "Difícil"
            else:
                readability_score = 30
                reading_level = "Muito Difícil"
            
            result = {
                "readability_score": readability_score,
                "reading_level": reading_level,
                "average_sentence_length": avg_sentence_length,
                "average_word_length": stats["average_characters_per_word"],
                "total_words": words,
                "total_sentences": sentences,
                "recommendations": []
            }
            
            # Add recommendations
            if avg_sentence_length > 20:
                result["recommendations"].append("Reduza o tamanho médio das frases para melhorar legibilidade")
            if stats["average_characters_per_word"] > 6:
                result["recommendations"].append("Use palavras mais curtas e simples quando possível")
            if readability_score < 50:
                result["recommendations"].append("Considere simplificar o texto para torná-lo mais acessível")
            
            # Use AI for qualitative analysis if available
            if self.gemini_client:
                try:
                    prompt = f"""Analise a legibilidade do seguinte texto:

Texto:
{text[:2000]}

Métricas:
- Tamanho médio de frase: {avg_sentence_length} palavras
- Pontuação de legibilidade: {readability_score}/100
- Nível de leitura: {reading_level}

Forneça uma análise qualitativa da legibilidade e sugestões específicas para melhorar."""
                    
                    ai_analysis = self.gemini_client.generate_response(prompt, timeout=60)
                    result["ai_analysis"] = ai_analysis
                except Exception as e:
                    logger.warning(f"AI analysis failed: {str(e)}")
            
            logger.info(f"Readability analysis complete: score={readability_score}")
            return result
            
        except Exception as e:
            logger.error(f"Error analyzing readability: {str(e)}")
            raise
    
    def check_term_consistency(self, file_path: str) -> Dict[str, Any]:
        """Verifica consistência de termos no documento.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with term consistency analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Checking term consistency in: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            log_file_access_error(logger, file_path, "arquivo não encontrado")
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text.strip():
                raise ValueError("Documento está vazio")
            
            # Extract words (simple tokenization)
            import re
            words = re.findall(r'\b[a-záàâãéèêíïóôõöúçñA-ZÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ]+\b', text.lower())
            
            # Count word frequencies
            from collections import Counter
            word_freq = Counter(words)
            
            # Find potential inconsistencies (words that are very similar)
            # This is a simplified approach - in production would use more sophisticated methods
            potential_variations = {}
            checked = set()
            
            for word1 in word_freq:
                if word1 in checked or len(word1) < 4:
                    continue
                
                variations = [word1]
                for word2 in word_freq:
                    if word2 != word1 and word2 not in checked:
                        # Simple similarity check (same length, differ by 1-2 chars)
                        if abs(len(word1) - len(word2)) <= 1:
                            diff_count = sum(c1 != c2 for c1, c2 in zip(word1, word2))
                            if diff_count <= 2:
                                variations.append(word2)
                                checked.add(word2)
                
                if len(variations) > 1:
                    potential_variations[word1] = {
                        "variations": variations,
                        "occurrences": {v: word_freq[v] for v in variations}
                    }
                checked.add(word1)
            
            # Calculate consistency score
            total_words = len(words)
            inconsistent_words = sum(sum(v["occurrences"].values()) for v in potential_variations.values())
            consistency_score = round(1 - (inconsistent_words / total_words), 2) if total_words > 0 else 1.0
            
            result = {
                "consistency_score": consistency_score,
                "inconsistencies_found": len(potential_variations),
                "term_variations": [
                    {
                        "canonical": canonical,
                        "variations": data["variations"],
                        "occurrences": data["occurrences"]
                    }
                    for canonical, data in list(potential_variations.items())[:10]  # Top 10
                ],
                "total_unique_words": len(word_freq),
                "recommendations": []
            }
            
            if len(potential_variations) > 0:
                result["recommendations"].append(f"Padronize {len(potential_variations)} termo(s) com variações")
            if consistency_score < 0.9:
                result["recommendations"].append("Revise o documento para melhorar consistência terminológica")
            
            logger.info(f"Term consistency check complete: score={consistency_score}")
            return result
            
        except Exception as e:
            logger.error(f"Error checking term consistency: {str(e)}")
            raise
    
    def analyze_document(self, file_path: str, include_ai_analysis: bool = True) -> Dict[str, Any]:
        """Realiza análise completa do documento.
        
        Agrega todas as análises em um relatório único.
        
        Args:
            file_path: Path to the Word file
            include_ai_analysis: Whether to include AI-powered analyses (default: True)
            
        Returns:
            Dictionary with complete document analysis
            
        Raises:
            FileNotFoundError: If the file does not exist
        """
        logger.info(f"Performing complete document analysis: {file_path}")
        
        try:
            result = {
                "file_path": file_path,
                "statistics": self.get_document_statistics(file_path),
                "word_count": self.analyze_word_count(file_path),
                "section_length": self.analyze_section_length(file_path),
                "readability": self.analyze_readability(file_path),
                "term_consistency": self.check_term_consistency(file_path)
            }
            
            # Add AI-powered analyses if requested and available
            if include_ai_analysis and self.gemini_client:
                try:
                    result["tone"] = self.analyze_tone(file_path)
                    result["jargon"] = self.identify_jargon(file_path)
                except Exception as e:
                    logger.warning(f"AI analysis failed: {str(e)}")
                    result["ai_analysis_error"] = str(e)
            
            logger.info("Complete document analysis finished")
            return result
            
        except Exception as e:
            logger.error(f"Error in complete document analysis: {str(e)}")
            raise
