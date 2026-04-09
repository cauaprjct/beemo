"""Tests for AI-powered Word text improvement operations.

This module tests the new text improvement features that use Gemini AI
to correct grammar, improve clarity, adjust tone, simplify language, and rewrite text.
"""

import os
import pytest
from pathlib import Path
from docx import Document
from unittest.mock import Mock, MagicMock

from src.word_tool import WordTool
from src.exceptions import CorruptedFileError


@pytest.fixture
def temp_dir(tmp_path):
    """Create a temporary directory for test files."""
    return tmp_path


@pytest.fixture
def mock_gemini_client():
    """Create a mock Gemini client for testing."""
    client = Mock()
    client.generate_response = MagicMock()
    return client


@pytest.fixture
def word_tool_with_ai(mock_gemini_client):
    """Create WordTool instance with mocked Gemini client."""
    return WordTool(gemini_client=mock_gemini_client)


@pytest.fixture
def word_tool_no_ai():
    """Create WordTool instance without Gemini client."""
    return WordTool()


@pytest.fixture
def sample_document(temp_dir):
    """Create a sample Word document for testing."""
    file_path = os.path.join(temp_dir, "test_doc.docx")
    doc = Document()
    doc.add_paragraph("This is a test paragraph with some erors.")
    doc.add_paragraph("Another paragraph that could be improved.")
    doc.add_paragraph("A third paragraph for testing purposes.")
    doc.save(file_path)
    return file_path


class TestCorrectGrammar:
    """Tests for grammar correction functionality."""
    
    def test_correct_grammar_document(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test correcting grammar in entire document."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """This is a test paragraph with some errors.
Another paragraph that could be improved.
A third paragraph for testing purposes."""
        
        result = word_tool_with_ai.correct_grammar(sample_document, target="document")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert "errors" in result
        assert len(result) > 0
        
        # Verify document was updated
        doc = Document(sample_document)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        assert "errors" in text
    
    def test_correct_grammar_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test correcting grammar in specific paragraph."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "This is a test paragraph with some errors."
        
        result = word_tool_with_ai.correct_grammar(sample_document, target="paragraph", index=0)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert "errors" in result
        
        # Verify only first paragraph was updated
        doc = Document(sample_document)
        assert "errors" in doc.paragraphs[0].text
        # Note: Second paragraph should remain unchanged, but the original typo "erors" 
        # was in the first paragraph, not the second
    
    def test_correct_grammar_without_ai_client(self, word_tool_no_ai, sample_document):
        """Test that grammar correction fails without AI client."""
        with pytest.raises(ValueError, match="GeminiClient não configurado"):
            word_tool_no_ai.correct_grammar(sample_document)
    
    def test_correct_grammar_invalid_index(self, word_tool_with_ai, sample_document):
        """Test grammar correction with invalid paragraph index."""
        with pytest.raises(IndexError):
            word_tool_with_ai.correct_grammar(sample_document, target="paragraph", index=999)
    
    def test_correct_grammar_nonexistent_file(self, word_tool_with_ai, temp_dir):
        """Test grammar correction on non-existent file."""
        file_path = os.path.join(temp_dir, "nonexistent.docx")
        with pytest.raises(FileNotFoundError):
            word_tool_with_ai.correct_grammar(file_path)


class TestImproveClarity:
    """Tests for clarity improvement functionality."""
    
    def test_improve_clarity_document(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test improving clarity of entire document."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """This is a clear test paragraph.
Another well-structured paragraph.
A third concise paragraph."""
        
        result = word_tool_with_ai.improve_clarity(sample_document, target="document")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert "clear" in result or "well-structured" in result
        assert len(result) > 0
    
    def test_improve_clarity_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test improving clarity of specific paragraph."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "This is a clear and concise test paragraph."
        
        result = word_tool_with_ai.improve_clarity(sample_document, target="paragraph", index=0)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert "clear" in result or "concise" in result


class TestAdjustTone:
    """Tests for tone adjustment functionality."""
    
    def test_adjust_tone_formal(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test adjusting tone to formal."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """This document presents a test paragraph.
The following section provides additional information.
The final section concludes the document."""
        
        result = word_tool_with_ai.adjust_tone(sample_document, tone="formal", target="document")
        
        # Verify AI was called with tone parameter
        assert mock_gemini_client.generate_response.called
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "formal" in call_args.lower()
        
        # Verify result
        assert len(result) > 0
    
    def test_adjust_tone_informal(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test adjusting tone to informal."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """Hey, this is a test paragraph.
Here's another paragraph.
And here's the last one."""
        
        result = word_tool_with_ai.adjust_tone(sample_document, tone="informal", target="document")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
    
    def test_adjust_tone_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test adjusting tone of specific paragraph."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "This document presents a formal test paragraph."
        
        result = word_tool_with_ai.adjust_tone(sample_document, tone="formal", target="paragraph", index=0)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0


class TestSimplifyLanguage:
    """Tests for language simplification functionality."""
    
    def test_simplify_language_document(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test simplifying language of entire document."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """This is a test paragraph.
Another paragraph.
A third paragraph."""
        
        result = word_tool_with_ai.simplify_language(sample_document, target="document")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
    
    def test_simplify_language_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test simplifying language of specific paragraph."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "This is a simple test paragraph."
        
        result = word_tool_with_ai.simplify_language(sample_document, target="paragraph", index=0)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert "simple" in result


class TestRewriteProfessional:
    """Tests for professional rewriting functionality."""
    
    def test_rewrite_professional_document(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test rewriting entire document professionally."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """This document presents a comprehensive test paragraph with professional formatting.
The subsequent section provides additional context and relevant information.
The concluding section summarizes the key points discussed."""
        
        result = word_tool_with_ai.rewrite_professional(sample_document, target="document")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
    
    def test_rewrite_professional_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test rewriting specific paragraph professionally."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "This document presents a professionally formatted test paragraph."
        
        result = word_tool_with_ai.rewrite_professional(sample_document, target="paragraph", index=0)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0


class TestImproveTextGeneric:
    """Tests for generic improve_text method."""
    
    def test_improve_text_with_all_types(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test improve_text with different improvement types."""
        improvement_types = ['grammar', 'clarity', 'tone', 'simplify', 'rewrite']
        
        for imp_type in improvement_types:
            # Mock AI response
            mock_gemini_client.generate_response.return_value = "Improved text."
            
            result = word_tool_with_ai.improve_text(
                sample_document, 
                improvement_type=imp_type,
                target="paragraph",
                index=0,
                tone="formal" if imp_type == 'tone' else None
            )
            
            # Verify AI was called
            assert mock_gemini_client.generate_response.called
            assert len(result) > 0
            
            # Reset mock
            mock_gemini_client.reset_mock()
    
    def test_improve_text_invalid_target(self, word_tool_with_ai, sample_document):
        """Test improve_text with invalid target."""
        with pytest.raises(ValueError, match="Target inválido"):
            word_tool_with_ai.improve_text(sample_document, 'grammar', target="invalid")
    
    def test_improve_text_empty_document(self, word_tool_with_ai, temp_dir, mock_gemini_client):
        """Test improve_text on empty document."""
        file_path = os.path.join(temp_dir, "empty.docx")
        doc = Document()
        doc.add_paragraph("")  # Empty paragraph
        doc.save(file_path)
        
        with pytest.raises(ValueError, match="vazio"):
            word_tool_with_ai.improve_text(file_path, 'grammar', target="document")


class TestResponseCleaning:
    """Tests for AI response cleaning functionality."""
    
    def test_clean_ai_response_with_markdown(self, word_tool_with_ai):
        """Test cleaning AI response with markdown code blocks."""
        response = "```\nThis is the improved text.\n```"
        cleaned = word_tool_with_ai._clean_ai_response(response)
        assert "```" not in cleaned
        assert "This is the improved text." in cleaned
    
    def test_clean_ai_response_with_prefixes(self, word_tool_with_ai):
        """Test cleaning AI response with common prefixes."""
        response = "TEXTO CORRIGIDO: This is the corrected text."
        cleaned = word_tool_with_ai._clean_ai_response(response)
        assert "TEXTO CORRIGIDO:" not in cleaned
        assert "This is the corrected text." in cleaned
    
    def test_clean_ai_response_with_whitespace(self, word_tool_with_ai):
        """Test cleaning AI response with extra whitespace."""
        response = "\n\n  This is the text.  \n\n"
        cleaned = word_tool_with_ai._clean_ai_response(response)
        assert cleaned == "This is the text."


class TestPromptBuilding:
    """Tests for improvement prompt building."""
    
    def test_build_improvement_prompt_grammar(self, word_tool_with_ai):
        """Test building grammar correction prompt."""
        prompt = word_tool_with_ai._build_improvement_prompt("Test text", "grammar")
        assert "gramaticais" in prompt.lower() or "gramática" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_improvement_prompt_tone(self, word_tool_with_ai):
        """Test building tone adjustment prompt."""
        prompt = word_tool_with_ai._build_improvement_prompt("Test text", "tone", tone="formal")
        assert "formal" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_improvement_prompt_simplify(self, word_tool_with_ai):
        """Test building simplification prompt."""
        prompt = word_tool_with_ai._build_improvement_prompt("Test text", "simplify")
        assert "simplif" in prompt.lower()
        assert "Test text" in prompt


class TestEdgeCases:
    """Tests for edge cases and error handling."""
    
    def test_improve_text_with_very_long_text(self, word_tool_with_ai, temp_dir, mock_gemini_client):
        """Test improving very long text (performance consideration)."""
        file_path = os.path.join(temp_dir, "long_doc.docx")
        doc = Document()
        
        # Create document with many paragraphs
        for i in range(50):
            doc.add_paragraph(f"This is paragraph number {i} with some content.")
        doc.save(file_path)
        
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "Improved text." * 50
        
        result = word_tool_with_ai.improve_text(file_path, 'grammar', target="document")
        
        # Verify it completes without timeout
        assert len(result) > 0
    
    def test_improve_text_ai_returns_empty(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test handling when AI returns empty or invalid response."""
        # Mock AI returning empty response
        mock_gemini_client.generate_response.return_value = ""
        
        with pytest.raises(ValueError, match="inválido ou muito curto"):
            word_tool_with_ai.improve_text(sample_document, 'grammar', target="document")
    
    def test_improve_text_ai_returns_short_text(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test handling when AI returns very short response."""
        # Mock AI returning very short response
        mock_gemini_client.generate_response.return_value = "Hi"
        
        with pytest.raises(ValueError, match="inválido ou muito curto"):
            word_tool_with_ai.improve_text(sample_document, 'grammar', target="document")
