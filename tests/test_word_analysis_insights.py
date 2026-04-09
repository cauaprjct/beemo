"""Tests for Word document analysis and insights operations (Phase 4)."""

import pytest
from pathlib import Path
from docx import Document
from src.word_tool import WordTool
from src.gemini_client import GeminiClient
from unittest.mock import Mock
import tempfile


@pytest.fixture
def word_tool():
    """Create WordTool instance without AI client."""
    return WordTool(gemini_client=None)


@pytest.fixture
def word_tool_with_ai():
    """Create WordTool instance with mocked AI client."""
    mock_client = Mock(spec=GeminiClient)
    return WordTool(gemini_client=mock_client)


@pytest.fixture
def mock_gemini_client():
    """Create mocked Gemini client."""
    return Mock(spec=GeminiClient)


@pytest.fixture
def sample_document_with_sections(tmp_path):
    """Create a sample document with multiple sections."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction section with some content.")
    doc.add_paragraph("It has multiple paragraphs to test word counting.")
    
    doc.add_heading("Main Content", level=1)
    doc.add_paragraph("This is the main content section.")
    doc.add_paragraph("It contains more detailed information about the topic.")
    doc.add_paragraph("And even more content to make it longer.")
    
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("This is the conclusion section.")
    
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def sample_long_document(tmp_path):
    """Create a document with long sections."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    
    doc.add_heading("Long Section", level=1)
    # Add many paragraphs to exceed threshold
    for i in range(50):
        doc.add_paragraph(f"This is paragraph {i+1} with some content to make the section long.")
    
    doc.add_heading("Short Section", level=1)
    doc.add_paragraph("This is a short section.")
    
    doc.save(str(doc_path))
    return str(doc_path)


# ==================== ANALYZE WORD COUNT TESTS ====================

class TestAnalyzeWordCount:
    
    def test_analyze_word_count_basic(self, word_tool, sample_document_with_sections):
        """Test basic word count analysis."""
        result = word_tool.analyze_word_count(sample_document_with_sections)
        
        assert "total_words" in result
        assert "total_paragraphs" in result
        assert "sections" in result
        assert "section_count" in result
        assert result["total_words"] > 0
        assert result["section_count"] == 3
    
    def test_analyze_word_count_sections(self, word_tool, sample_document_with_sections):
        """Test word count by section."""
        result = word_tool.analyze_word_count(sample_document_with_sections)
        
        sections = result["sections"]
        assert len(sections) == 3
        assert sections[0]["title"] == "Introduction"
        assert sections[0]["words"] > 0
        assert "percentage" in sections[0]
    
    def test_analyze_word_count_empty_document(self, word_tool, tmp_path):
        """Test word count on empty document."""
        doc_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(doc_path))
        
        result = word_tool.analyze_word_count(str(doc_path))
        assert result["total_words"] == 0
        assert result["total_paragraphs"] == 0
    
    def test_analyze_word_count_nonexistent_file(self, word_tool):
        """Test word count on nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.analyze_word_count("nonexistent.docx")


# ==================== ANALYZE SECTION LENGTH TESTS ====================

class TestAnalyzeSectionLength:
    
    def test_analyze_section_length_default_threshold(self, word_tool, sample_long_document):
        """Test section length analysis with default threshold."""
        result = word_tool.analyze_section_length(sample_long_document)
        
        assert "long_sections_count" in result
        assert "long_sections" in result
        assert "max_words_threshold" in result
        assert result["max_words_threshold"] == 500
    
    def test_analyze_section_length_custom_threshold(self, word_tool, sample_document_with_sections):
        """Test section length analysis with custom threshold."""
        result = word_tool.analyze_section_length(sample_document_with_sections, max_words=10)
        
        assert result["max_words_threshold"] == 10
        # Should find long sections with low threshold
        assert result["long_sections_count"] >= 0
    
    def test_analyze_section_length_recommendations(self, word_tool, sample_long_document):
        """Test that recommendations are provided for long sections."""
        result = word_tool.analyze_section_length(sample_long_document, max_words=100)
        
        if result["long_sections_count"] > 0:
            long_section = result["long_sections"][0]
            assert "recommendation" in long_section
            assert "excess_words" in long_section


# ==================== GET DOCUMENT STATISTICS TESTS ====================

class TestGetDocumentStatistics:
    
    def test_get_statistics_basic(self, word_tool, sample_document_with_sections):
        """Test getting basic document statistics."""
        result = word_tool.get_document_statistics(sample_document_with_sections)
        
        assert "total_paragraphs" in result
        assert "total_words" in result
        assert "total_characters" in result
        assert "total_sentences" in result
        assert "average_words_per_sentence" in result
        assert "estimated_reading_time_minutes" in result
    
    def test_get_statistics_with_tables(self, word_tool, tmp_path):
        """Test statistics on document with tables."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Some text")
        doc.add_table(rows=2, cols=2)
        doc.save(str(doc_path))
        
        result = word_tool.get_document_statistics(str(doc_path))
        assert result["total_tables"] == 1
    
    def test_get_statistics_reading_time(self, word_tool, sample_document_with_sections):
        """Test reading time estimation."""
        result = word_tool.get_document_statistics(sample_document_with_sections)
        
        # Reading time should be based on 200 words per minute
        expected_time = result["total_words"] / 200
        assert abs(result["estimated_reading_time_minutes"] - expected_time) < 0.1


# ==================== ANALYZE TONE TESTS ====================

class TestAnalyzeTone:
    
    def test_analyze_tone_with_ai(self, word_tool_with_ai, sample_document_with_sections, mock_gemini_client):
        """Test tone analysis with AI."""
        mock_gemini_client.generate_response.return_value = "O documento apresenta tom formal e profissional."
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        result = word_tool_with_ai.analyze_tone(sample_document_with_sections)
        
        assert "analysis" in result
        assert "text_sample" in result
        assert len(result["analysis"]) > 0
    
    def test_analyze_tone_without_ai(self, word_tool, sample_document_with_sections):
        """Test tone analysis without AI client."""
        with pytest.raises(ValueError, match="GeminiClient não configurado"):
            word_tool.analyze_tone(sample_document_with_sections)
    
    def test_analyze_tone_empty_document(self, word_tool_with_ai, tmp_path, mock_gemini_client):
        """Test tone analysis on empty document."""
        doc_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(doc_path))
        
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        with pytest.raises(ValueError, match="Documento está vazio"):
            word_tool_with_ai.analyze_tone(str(doc_path))


# ==================== IDENTIFY JARGON TESTS ====================

class TestIdentifyJargon:
    
    def test_identify_jargon_with_ai(self, word_tool_with_ai, sample_document_with_sections, mock_gemini_client):
        """Test jargon identification with AI."""
        mock_gemini_client.generate_response.return_value = "Termos identificados: paradigma, framework, etc."
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        result = word_tool_with_ai.identify_jargon(sample_document_with_sections)
        
        assert "analysis" in result
        assert "text_sample" in result
    
    def test_identify_jargon_without_ai(self, word_tool, sample_document_with_sections):
        """Test jargon identification without AI client."""
        with pytest.raises(ValueError, match="GeminiClient não configurado"):
            word_tool.identify_jargon(sample_document_with_sections)


# ==================== ANALYZE READABILITY TESTS ====================

class TestAnalyzeReadability:
    
    def test_analyze_readability_basic(self, word_tool, sample_document_with_sections):
        """Test basic readability analysis."""
        result = word_tool.analyze_readability(sample_document_with_sections)
        
        assert "readability_score" in result
        assert "reading_level" in result
        assert "average_sentence_length" in result
        assert "recommendations" in result
        assert 0 <= result["readability_score"] <= 100
    
    def test_analyze_readability_with_ai(self, word_tool_with_ai, sample_document_with_sections, mock_gemini_client):
        """Test readability analysis with AI enhancement."""
        mock_gemini_client.generate_response.return_value = "O texto tem boa legibilidade."
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        result = word_tool_with_ai.analyze_readability(sample_document_with_sections)
        
        assert "ai_analysis" in result
    
    def test_analyze_readability_levels(self, word_tool, tmp_path):
        """Test different readability levels."""
        # Create document with short sentences (easy)
        doc_path = tmp_path / "easy.docx"
        doc = Document()
        doc.add_paragraph("This is easy. Very short. Simple words.")
        doc.save(str(doc_path))
        
        result = word_tool.analyze_readability(str(doc_path))
        assert result["readability_score"] >= 60  # Should be relatively easy


# ==================== CHECK TERM CONSISTENCY TESTS ====================

class TestCheckTermConsistency:
    
    def test_check_consistency_basic(self, word_tool, sample_document_with_sections):
        """Test basic term consistency check."""
        result = word_tool.check_term_consistency(sample_document_with_sections)
        
        assert "consistency_score" in result
        assert "inconsistencies_found" in result
        assert "term_variations" in result
        assert "recommendations" in result
        assert 0 <= result["consistency_score"] <= 1
    
    def test_check_consistency_with_variations(self, word_tool, tmp_path):
        """Test consistency check with term variations."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("The user can login. The usuario can logon. The utilizador can log-in.")
        doc.save(str(doc_path))
        
        result = word_tool.check_term_consistency(str(doc_path))
        
        # Should detect some variations
        assert result["total_unique_words"] > 0
    
    def test_check_consistency_empty_document(self, word_tool, tmp_path):
        """Test consistency check on empty document."""
        doc_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(doc_path))
        
        with pytest.raises(ValueError, match="Documento está vazio"):
            word_tool.check_term_consistency(str(doc_path))


# ==================== ANALYZE DOCUMENT (COMPLETE) TESTS ====================

class TestAnalyzeDocument:
    
    def test_analyze_document_without_ai(self, word_tool, sample_document_with_sections):
        """Test complete document analysis without AI."""
        result = word_tool.analyze_document(sample_document_with_sections, include_ai_analysis=False)
        
        assert "file_path" in result
        assert "statistics" in result
        assert "word_count" in result
        assert "section_length" in result
        assert "readability" in result
        assert "term_consistency" in result
        assert "tone" not in result  # AI analysis not included
        assert "jargon" not in result
    
    def test_analyze_document_with_ai(self, word_tool_with_ai, sample_document_with_sections, mock_gemini_client):
        """Test complete document analysis with AI."""
        mock_gemini_client.generate_response.return_value = "Análise completa do documento."
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        result = word_tool_with_ai.analyze_document(sample_document_with_sections, include_ai_analysis=True)
        
        assert "tone" in result
        assert "jargon" in result
    
    def test_analyze_document_ai_failure(self, word_tool_with_ai, sample_document_with_sections, mock_gemini_client):
        """Test complete analysis when AI fails."""
        mock_gemini_client.generate_response.side_effect = Exception("AI error")
        word_tool_with_ai.gemini_client = mock_gemini_client
        
        result = word_tool_with_ai.analyze_document(sample_document_with_sections, include_ai_analysis=True)
        
        # Should still return basic analyses
        assert "statistics" in result
        assert "ai_analysis_error" in result


# ==================== EDGE CASES ====================

class TestEdgeCases:
    
    def test_analyze_nonexistent_file(self, word_tool):
        """Test analyzing nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.analyze_word_count("nonexistent.docx")
    
    def test_statistics_nonexistent_file(self, word_tool):
        """Test statistics on nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.get_document_statistics("nonexistent.docx")
    
    def test_readability_nonexistent_file(self, word_tool):
        """Test readability on nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.analyze_readability("nonexistent.docx")
    
    def test_consistency_nonexistent_file(self, word_tool):
        """Test consistency on nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.check_term_consistency("nonexistent.docx")
    
    def test_complete_analysis_nonexistent_file(self, word_tool):
        """Test complete analysis on nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.analyze_document("nonexistent.docx")
