"""Unit tests for WordTool class.

Tests the Word file manipulation functionality including reading, creating,
updating, and extracting data from Word documents.
"""

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document

from src.exceptions import CorruptedFileError
from src.word_tool import WordTool


@pytest.fixture
def word_tool():
    """Fixture providing a WordTool instance."""
    return WordTool()


@pytest.fixture
def temp_dir():
    """Fixture providing a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def sample_word_file(temp_dir):
    """Fixture creating a sample Word file for testing."""
    file_path = os.path.join(temp_dir, "sample.docx")
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    doc.save(file_path)
    return file_path


@pytest.fixture
def word_file_with_table(temp_dir):
    """Fixture creating a Word file with a table."""
    file_path = os.path.join(temp_dir, "with_table.docx")
    doc = Document()
    doc.add_paragraph("Document with table")
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f"Cell {i},{j}"
    
    doc.save(file_path)
    return file_path


class TestReadWord:
    """Tests for read_word method."""
    
    def test_read_existing_file(self, word_tool, sample_word_file):
        """Test reading an existing Word file."""
        content = word_tool.read_word(sample_word_file)
        
        assert isinstance(content, str)
        assert "First paragraph" in content
        assert "Second paragraph" in content
        assert "Third paragraph" in content
    
    def test_read_nonexistent_file(self, word_tool, temp_dir):
        """Test reading a non-existent file raises FileNotFoundError."""
        file_path = os.path.join(temp_dir, "nonexistent.docx")
        
        with pytest.raises(FileNotFoundError) as exc_info:
            word_tool.read_word(file_path)
        
        assert "Word file not found" in str(exc_info.value)
    
    def test_read_corrupted_file(self, word_tool, temp_dir):
        """Test reading a corrupted file raises CorruptedFileError."""
        file_path = os.path.join(temp_dir, "corrupted.docx")
        
        # Create a corrupted file (not a valid docx)
        with open(file_path, "w") as f:
            f.write("This is not a valid Word document")
        
        with pytest.raises(CorruptedFileError) as exc_info:
            word_tool.read_word(file_path)
        
        assert "corrupted or invalid" in str(exc_info.value)


class TestCreateWord:
    """Tests for create_word method."""
    
    def test_create_simple_document(self, word_tool, temp_dir):
        """Test creating a simple Word document."""
        file_path = os.path.join(temp_dir, "new_doc.docx")
        content = "Hello World\nThis is a test"
        
        word_tool.create_word(file_path, content)
        
        # Verify file was created
        assert os.path.exists(file_path)
        
        # Verify content
        result = word_tool.read_word(file_path)
        assert "Hello World" in result
        assert "This is a test" in result
    
    def test_create_empty_document(self, word_tool, temp_dir):
        """Test creating an empty Word document."""
        file_path = os.path.join(temp_dir, "empty_doc.docx")
        
        word_tool.create_word(file_path, "")
        
        assert os.path.exists(file_path)


class TestAddParagraph:
    """Tests for add_paragraph method."""
    
    def test_add_paragraph_to_existing_file(self, word_tool, sample_word_file):
        """Test adding a paragraph to an existing file."""
        new_text = "Fourth paragraph"
        
        word_tool.add_paragraph(sample_word_file, new_text)
        
        # Verify paragraph was added
        content = word_tool.read_word(sample_word_file)
        assert new_text in content
    
    def test_add_paragraph_to_nonexistent_file(self, word_tool, temp_dir):
        """Test adding paragraph to non-existent file raises FileNotFoundError."""
        file_path = os.path.join(temp_dir, "nonexistent.docx")
        
        with pytest.raises(FileNotFoundError):
            word_tool.add_paragraph(file_path, "Some text")


class TestUpdateParagraph:
    """Tests for update_paragraph method."""
    
    def test_update_existing_paragraph(self, word_tool, sample_word_file):
        """Test updating an existing paragraph."""
        new_text = "Updated first paragraph"
        
        word_tool.update_paragraph(sample_word_file, 0, new_text)
        
        # Verify paragraph was updated
        content = word_tool.read_word(sample_word_file)
        assert new_text in content
        assert "First paragraph" not in content
    
    def test_update_paragraph_invalid_index(self, word_tool, sample_word_file):
        """Test updating with invalid index raises IndexError."""
        with pytest.raises(IndexError) as exc_info:
            word_tool.update_paragraph(sample_word_file, 999, "New text")
        
        assert "out of range" in str(exc_info.value)
    
    def test_update_paragraph_negative_index(self, word_tool, sample_word_file):
        """Test updating with negative index raises IndexError."""
        with pytest.raises(IndexError):
            word_tool.update_paragraph(sample_word_file, -1, "New text")


class TestExtractTables:
    """Tests for extract_tables method."""
    
    def test_extract_tables_from_document(self, word_tool, word_file_with_table):
        """Test extracting tables from a document."""
        tables = word_tool.extract_tables(word_file_with_table)
        
        assert isinstance(tables, list)
        assert len(tables) == 1
        assert len(tables[0]) == 3  # 3 rows
        assert len(tables[0][0]) == 3  # 3 columns
        assert tables[0][0][0] == "Cell 0,0"
    
    def test_extract_tables_from_document_without_tables(self, word_tool, sample_word_file):
        """Test extracting tables from a document without tables."""
        tables = word_tool.extract_tables(sample_word_file)
        
        assert isinstance(tables, list)
        assert len(tables) == 0
    
    def test_extract_tables_nonexistent_file(self, word_tool, temp_dir):
        """Test extracting tables from non-existent file raises FileNotFoundError."""
        file_path = os.path.join(temp_dir, "nonexistent.docx")
        
        with pytest.raises(FileNotFoundError):
            word_tool.extract_tables(file_path)


class TestRoundTrip:
    """Tests for round-trip operations (create and read back)."""
    
    def test_create_and_read_round_trip(self, word_tool, temp_dir):
        """Test that creating and reading a document preserves content."""
        file_path = os.path.join(temp_dir, "roundtrip.docx")
        original_content = "Line 1\nLine 2\nLine 3"
        
        word_tool.create_word(file_path, original_content)
        read_content = word_tool.read_word(file_path)
        
        # Verify all lines are present
        for line in original_content.split("\n"):
            assert line in read_content
