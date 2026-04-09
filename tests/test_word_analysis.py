"""Tests for AI-powered Word content analysis and generation operations.

This module tests the new content analysis features that use Gemini AI
to generate summaries, extract key points, create resumes, generate conclusions, and create FAQs.
"""

import os
import pytest
from pathlib import Path
from docx import Document
from unittest.mock import Mock, MagicMock

from src.word_tool import WordTool
from src.exceptions import CorruptedFileError


@pytest.fixture
def temp_dir(tmp_path):
    """Create a temporary directory for test files."""
    return tmp_path


@pytest.fixture
def mock_gemini_client():
    """Create a mock Gemini client for testing."""
    client = Mock()
    client.generate_response = MagicMock()
    return client


@pytest.fixture
def word_tool_with_ai(mock_gemini_client):
    """Create WordTool instance with mocked Gemini client."""
    return WordTool(gemini_client=mock_gemini_client)


@pytest.fixture
def word_tool_no_ai():
    """Create WordTool instance without Gemini client."""
    return WordTool()


@pytest.fixture
def sample_document(temp_dir):
    """Create a sample Word document for testing."""
    file_path = os.path.join(temp_dir, "test_doc.docx")
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a comprehensive document about project management.")
    doc.add_paragraph("It covers planning, execution, and monitoring phases.")
    doc.add_heading("Planning Phase", level=1)
    doc.add_paragraph("Planning involves defining objectives and creating schedules.")
    doc.add_paragraph("Resource allocation is a critical part of planning.")
    doc.add_heading("Execution Phase", level=1)
    doc.add_paragraph("Execution is where the actual work happens.")
    doc.add_paragraph("Team coordination is essential during execution.")
    doc.save(file_path)
    return file_path


class TestGenerateSummary:
    """Tests for summary generation functionality."""
    
    def test_generate_summary_new_section(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test generating summary as new section."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """Este documento apresenta uma visão abrangente sobre gerenciamento de projetos.
Aborda três fases principais: planejamento, execução e monitoramento.
O planejamento envolve definição de objetivos e criação de cronogramas.
A execução é onde o trabalho real acontece, com coordenação de equipe sendo essencial."""
        
        result = word_tool_with_ai.generate_summary(sample_document, output_mode="new_section")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
        assert "gerenciamento" in result.lower() or "projeto" in result.lower()
        
        # Verify document was updated with new section
        doc = Document(sample_document)
        paragraphs = [p.text for p in doc.paragraphs]
        # Should have original content + heading + summary paragraphs
        assert len(paragraphs) > 8  # Original had 8 paragraphs
    
    def test_generate_summary_custom_title(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test generating summary with custom section title."""
        mock_gemini_client.generate_response.return_value = "Summary content here."
        
        result = word_tool_with_ai.generate_summary(
            sample_document, 
            output_mode="new_section",
            section_title="Executive Summary"
        )
        
        assert mock_gemini_client.generate_response.called
        assert len(result) > 0
    
    def test_generate_summary_append_mode(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test generating summary in append mode."""
        mock_gemini_client.generate_response.return_value = "This is a comprehensive summary of the document content with sufficient length."
        
        result = word_tool_with_ai.generate_summary(sample_document, output_mode="append")
        
        assert mock_gemini_client.generate_response.called
        assert len(result) > 0
    
    def test_generate_summary_without_ai_client(self, word_tool_no_ai, sample_document):
        """Test that summary generation fails without AI client."""
        with pytest.raises(ValueError, match="GeminiClient não configurado"):
            word_tool_no_ai.generate_summary(sample_document)


class TestExtractKeyPoints:
    """Tests for key points extraction functionality."""
    
    def test_extract_key_points_default(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test extracting default number of key points."""
        # Mock AI response
        mock_gemini_client.generate_response.return_value = """1. Gerenciamento de projetos envolve três fases principais
2. Planejamento define objetivos e cronogramas
3. Alocação de recursos é crítica no planejamento
4. Execução é onde o trabalho real acontece
5. Coordenação de equipe é essencial na execução"""
        
        result = word_tool_with_ai.extract_key_points(sample_document, num_points=5)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
        
        # Verify document was updated with numbered list
        doc = Document(sample_document)
        # Check that list items were added
        assert len(doc.paragraphs) > 8
    
    def test_extract_key_points_custom_number(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test extracting custom number of key points."""
        mock_gemini_client.generate_response.return_value = """1. Point one
2. Point two
3. Point three"""
        
        result = word_tool_with_ai.extract_key_points(sample_document, num_points=3)
        
        assert mock_gemini_client.generate_response.called
        # Verify prompt includes the number
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "3" in call_args


class TestCreateResume:
    """Tests for resume creation functionality."""
    
    def test_create_resume_one_paragraph(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test creating 1-paragraph resume."""
        mock_gemini_client.generate_response.return_value = "Este documento aborda gerenciamento de projetos em três fases: planejamento, execução e monitoramento, com foco em objetivos, recursos e coordenação de equipe."
        
        result = word_tool_with_ai.create_resume(sample_document, size="1_paragraph")
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
        
        # Verify prompt includes size instruction
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "parágrafo" in call_args.lower()
    
    def test_create_resume_three_sentences(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test creating 3-sentence resume."""
        mock_gemini_client.generate_response.return_value = "Documento sobre gerenciamento de projetos. Cobre planejamento e execução. Foca em coordenação e recursos."
        
        result = word_tool_with_ai.create_resume(sample_document, size="3_sentences")
        
        assert mock_gemini_client.generate_response.called
        
        # Verify prompt includes size instruction
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "3 frases" in call_args.lower()
    
    def test_create_resume_one_page(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test creating 1-page resume."""
        mock_gemini_client.generate_response.return_value = "Long resume content here..." * 20
        
        result = word_tool_with_ai.create_resume(sample_document, size="1_page")
        
        assert mock_gemini_client.generate_response.called
        
        # Verify prompt includes size instruction
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "página" in call_args.lower()


class TestGenerateConclusions:
    """Tests for conclusions generation functionality."""
    
    def test_generate_conclusions_default(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test generating default number of conclusions."""
        mock_gemini_client.generate_response.return_value = """1. Planejamento adequado é fundamental para o sucesso do projeto
2. Coordenação de equipe impacta diretamente a execução
3. Alocação eficiente de recursos otimiza resultados"""
        
        result = word_tool_with_ai.generate_conclusions(sample_document, num_conclusions=3)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
        
        # Verify document was updated
        doc = Document(sample_document)
        assert len(doc.paragraphs) > 8
    
    def test_generate_conclusions_custom_number(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test generating custom number of conclusions."""
        mock_gemini_client.generate_response.return_value = """1. Conclusion one
2. Conclusion two
3. Conclusion three
4. Conclusion four
5. Conclusion five"""
        
        result = word_tool_with_ai.generate_conclusions(sample_document, num_conclusions=5)
        
        assert mock_gemini_client.generate_response.called
        
        # Verify prompt includes the number
        call_args = mock_gemini_client.generate_response.call_args[0][0]
        assert "5" in call_args


class TestCreateFAQ:
    """Tests for FAQ creation functionality."""
    
    def test_create_faq_default(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test creating default number of FAQ items."""
        mock_gemini_client.generate_response.return_value = """P: O que é gerenciamento de projetos?
R: É o processo de planejar, executar e monitorar projetos.

P: Quais são as fases principais?
R: Planejamento, execução e monitoramento.

P: Por que o planejamento é importante?
R: Define objetivos e cria cronogramas para o projeto."""
        
        result = word_tool_with_ai.create_faq(sample_document, num_questions=3)
        
        # Verify AI was called
        assert mock_gemini_client.generate_response.called
        
        # Verify result
        assert len(result) > 0
        assert "P:" in result or "Q:" in result
        assert "R:" in result or "A:" in result
        
        # Verify document was updated
        doc = Document(sample_document)
        assert len(doc.paragraphs) > 8
    
    def test_create_faq_custom_number(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test creating custom number of FAQ items."""
        mock_gemini_client.generate_response.return_value = """P: Question 1?
R: Answer 1.

P: Question 2?
R: Answer 2."""
        
        result = word_tool_with_ai.create_faq(sample_document, num_questions=2)
        
        assert mock_gemini_client.generate_response.called


class TestAnalyzeAndGenerate:
    """Tests for generic analyze_and_generate method."""
    
    def test_analyze_and_generate_all_types(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test analyze_and_generate with different analysis types."""
        analysis_types = ['summary', 'key_points', 'resume', 'conclusions', 'faq']
        
        for analysis_type in analysis_types:
            # Mock AI response
            mock_gemini_client.generate_response.return_value = "Generated content for " + analysis_type
            
            result = word_tool_with_ai.analyze_and_generate(
                sample_document,
                analysis_type=analysis_type,
                output_mode="append"
            )
            
            # Verify AI was called
            assert mock_gemini_client.generate_response.called
            assert len(result) > 0
            
            # Reset mock
            mock_gemini_client.reset_mock()
    
    def test_analyze_and_generate_invalid_output_mode(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test analyze_and_generate with invalid output mode."""
        mock_gemini_client.generate_response.return_value = "This is sufficient content for the test to pass validation."
        
        with pytest.raises(ValueError, match="Output mode inválido"):
            word_tool_with_ai.analyze_and_generate(
                sample_document,
                'summary',
                output_mode="invalid_mode"
            )
    
    def test_analyze_and_generate_empty_document(self, word_tool_with_ai, temp_dir, mock_gemini_client):
        """Test analyze_and_generate on empty document."""
        file_path = os.path.join(temp_dir, "empty.docx")
        doc = Document()
        doc.add_paragraph("")  # Empty paragraph
        doc.save(file_path)
        
        with pytest.raises(ValueError, match="vazio"):
            word_tool_with_ai.analyze_and_generate(file_path, 'summary')


class TestPromptBuilding:
    """Tests for analysis prompt building."""
    
    def test_build_analysis_prompt_summary(self, word_tool_with_ai):
        """Test building summary prompt."""
        prompt = word_tool_with_ai._build_analysis_prompt("Test text", "summary")
        assert "sumário" in prompt.lower() or "executivo" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_analysis_prompt_key_points(self, word_tool_with_ai):
        """Test building key points prompt."""
        prompt = word_tool_with_ai._build_analysis_prompt("Test text", "key_points", num_items=5)
        assert "5" in prompt
        assert "pontos" in prompt.lower() or "key" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_analysis_prompt_resume(self, word_tool_with_ai):
        """Test building resume prompt."""
        prompt = word_tool_with_ai._build_analysis_prompt("Test text", "resume", size="1_paragraph")
        assert "parágrafo" in prompt.lower() or "paragraph" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_analysis_prompt_conclusions(self, word_tool_with_ai):
        """Test building conclusions prompt."""
        prompt = word_tool_with_ai._build_analysis_prompt("Test text", "conclusions", num_items=3)
        assert "3" in prompt
        assert "conclus" in prompt.lower()
        assert "Test text" in prompt
    
    def test_build_analysis_prompt_faq(self, word_tool_with_ai):
        """Test building FAQ prompt."""
        prompt = word_tool_with_ai._build_analysis_prompt("Test text", "faq", num_items=5)
        assert "5" in prompt
        assert "faq" in prompt.lower() or "perguntas" in prompt.lower()
        assert "Test text" in prompt


class TestDefaultSectionTitles:
    """Tests for default section title generation."""
    
    def test_get_default_section_title(self, word_tool_with_ai):
        """Test getting default section titles for each analysis type."""
        titles = {
            'summary': 'Sumário Executivo',
            'key_points': 'Pontos-Chave',
            'resume': 'Resumo',
            'conclusions': 'Conclusões',
            'faq': 'Perguntas Frequentes (FAQ)'
        }
        
        for analysis_type, expected_title in titles.items():
            title = word_tool_with_ai._get_default_section_title(analysis_type)
            assert title == expected_title


class TestEdgeCases:
    """Tests for edge cases and error handling."""
    
    def test_analyze_nonexistent_file(self, word_tool_with_ai, temp_dir):
        """Test analyzing non-existent file."""
        file_path = os.path.join(temp_dir, "nonexistent.docx")
        with pytest.raises(FileNotFoundError):
            word_tool_with_ai.generate_summary(file_path)
    
    def test_analyze_ai_returns_empty(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test handling when AI returns empty response."""
        mock_gemini_client.generate_response.return_value = ""
        
        with pytest.raises(ValueError, match="inválido ou muito curto"):
            word_tool_with_ai.generate_summary(sample_document)
    
    def test_analyze_ai_returns_short_text(self, word_tool_with_ai, sample_document, mock_gemini_client):
        """Test handling when AI returns very short response."""
        mock_gemini_client.generate_response.return_value = "Short"
        
        with pytest.raises(ValueError, match="inválido ou muito curto"):
            word_tool_with_ai.generate_summary(sample_document)
    
    def test_analyze_with_very_long_document(self, word_tool_with_ai, temp_dir, mock_gemini_client):
        """Test analyzing very long document."""
        file_path = os.path.join(temp_dir, "long_doc.docx")
        doc = Document()
        
        # Create document with many paragraphs
        for i in range(100):
            doc.add_paragraph(f"This is paragraph number {i} with detailed content about the topic.")
        doc.save(file_path)
        
        # Mock AI response
        mock_gemini_client.generate_response.return_value = "Summary of the long document with key insights."
        
        result = word_tool_with_ai.generate_summary(file_path)
        
        # Verify it completes without timeout
        assert len(result) > 0
