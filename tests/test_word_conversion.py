"""Tests for Word format conversion and transformation operations (Phase 3)."""

import pytest
from pathlib import Path
from docx import Document
from src.word_tool import WordTool
from src.gemini_client import GeminiClient
from unittest.mock import Mock
import tempfile
import os


@pytest.fixture
def word_tool():
    """Create WordTool instance without AI client."""
    return WordTool(gemini_client=None)


@pytest.fixture
def sample_document_with_list(tmp_path):
    """Create a sample document with a numbered list."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    doc.add_paragraph("Introduction")
    doc.add_paragraph("1. First item", style='List Number')
    doc.add_paragraph("2. Second item", style='List Number')
    doc.add_paragraph("3. Third item", style='List Number')
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def sample_document_with_table(tmp_path):
    """Create a sample document with a table."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    doc.add_paragraph("Data Table")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Item 1"
    table.rows[1].cells[1].text = "100"
    table.rows[2].cells[0].text = "Item 2"
    table.rows[2].cells[1].text = "200"
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def sample_document_with_multiple_tables(tmp_path):
    """Create a sample document with multiple tables."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    
    # First table
    doc.add_paragraph("Sales Data")
    table1 = doc.add_table(rows=2, cols=2)
    table1.rows[0].cells[0].text = "Product"
    table1.rows[0].cells[1].text = "Sales"
    table1.rows[1].cells[0].text = "Widget"
    table1.rows[1].cells[1].text = "1000"
    
    # Second table
    doc.add_paragraph("Cost Data")
    table2 = doc.add_table(rows=2, cols=2)
    table2.rows[0].cells[0].text = "Item"
    table2.rows[0].cells[1].text = "Cost"
    table2.rows[1].cells[0].text = "Material"
    table2.rows[1].cells[1].text = "500"
    
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def sample_formatted_document(tmp_path):
    """Create a sample document with various formatting."""
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Section 1", level=2)
    doc.add_paragraph("This is a paragraph with some content.")
    doc.add_paragraph("1. First item", style='List Number')
    doc.add_paragraph("2. Second item", style='List Number')
    doc.add_heading("Section 2", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Data 1"
    table.rows[1].cells[1].text = "Data 2"
    doc.save(str(doc_path))
    return str(doc_path)


# ==================== CONVERT LIST TO TABLE TESTS ====================

class TestConvertListToTable:
    
    def test_convert_list_to_table_basic(self, word_tool, sample_document_with_list):
        """Test converting a numbered list to table."""
        result = word_tool.convert_list_to_table(sample_document_with_list, list_index=0)
        
        assert "convertida em tabela" in result
        assert "3 itens" in result
        
        # Verify table was created
        doc = Document(sample_document_with_list)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3
    
    def test_convert_list_to_table_with_header(self, word_tool, sample_document_with_list):
        """Test converting list to table with header."""
        result = word_tool.convert_list_to_table(
            sample_document_with_list,
            list_index=0,
            include_header=True,
            header_text="Tasks"
        )
        
        assert "convertida em tabela" in result
        
        # Verify header was added
        doc = Document(sample_document_with_list)
        assert len(doc.tables[0].rows) == 4  # 3 items + 1 header
        assert doc.tables[0].rows[0].cells[0].text == "Tasks"
    
    def test_convert_list_invalid_index(self, word_tool, sample_document_with_list):
        """Test converting list with invalid index."""
        with pytest.raises(ValueError, match="Índice de lista inválido"):
            word_tool.convert_list_to_table(sample_document_with_list, list_index=5)
    
    def test_convert_list_no_lists_found(self, word_tool, tmp_path):
        """Test converting when no lists exist."""
        doc_path = tmp_path / "no_lists.docx"
        doc = Document()
        doc.add_paragraph("Just a paragraph")
        doc.save(str(doc_path))
        
        with pytest.raises(ValueError, match="Nenhuma lista encontrada"):
            word_tool.convert_list_to_table(str(doc_path), list_index=0)


# ==================== CONVERT TABLE TO LIST TESTS ====================

class TestConvertTableToList:
    
    def test_convert_table_to_numbered_list(self, word_tool, sample_document_with_table):
        """Test converting table to numbered list."""
        result = word_tool.convert_table_to_list(
            sample_document_with_table,
            table_index=0,
            list_type="numbered",
            skip_header=True
        )
        
        assert "convertida em lista numbered" in result
        assert "2 itens" in result
        
        # Verify list was created
        doc = Document(sample_document_with_table)
        assert len(doc.tables) == 0  # Table should be removed
    
    def test_convert_table_to_bullet_list(self, word_tool, sample_document_with_table):
        """Test converting table to bullet list."""
        result = word_tool.convert_table_to_list(
            sample_document_with_table,
            table_index=0,
            list_type="bullet",
            skip_header=False
        )
        
        assert "convertida em lista bullet" in result
        assert "3 itens" in result  # Including header
    
    def test_convert_table_custom_separator(self, word_tool, sample_document_with_table):
        """Test converting table with custom separator."""
        result = word_tool.convert_table_to_list(
            sample_document_with_table,
            table_index=0,
            list_type="numbered",
            skip_header=True,
            separator=" | "
        )
        
        assert "convertida em lista" in result
    
    def test_convert_table_invalid_type(self, word_tool, sample_document_with_table):
        """Test converting table with invalid list type."""
        with pytest.raises(ValueError, match="Tipo de lista inválido"):
            word_tool.convert_table_to_list(
                sample_document_with_table,
                table_index=0,
                list_type="invalid"
            )
    
    def test_convert_table_invalid_index(self, word_tool, sample_document_with_table):
        """Test converting table with invalid index."""
        with pytest.raises(ValueError, match="Índice de tabela inválido"):
            word_tool.convert_table_to_list(sample_document_with_table, table_index=5)
    
    def test_convert_table_no_tables_found(self, word_tool, tmp_path):
        """Test converting when no tables exist."""
        doc_path = tmp_path / "no_tables.docx"
        doc = Document()
        doc.add_paragraph("Just a paragraph")
        doc.save(str(doc_path))
        
        with pytest.raises(ValueError, match="Nenhuma tabela encontrada"):
            word_tool.convert_table_to_list(str(doc_path), table_index=0)


# ==================== EXTRACT TABLES TO EXCEL TESTS ====================

class TestExtractTablesToExcel:
    
    def test_extract_tables_to_excel_basic(self, word_tool, sample_document_with_table, tmp_path):
        """Test extracting tables to Excel."""
        output_path = tmp_path / "output.xlsx"
        result = word_tool.extract_tables_to_excel(
            sample_document_with_table,
            str(output_path)
        )
        
        assert "Extraídas 1 tabela(s)" in result
        assert output_path.exists()
        
        # Verify Excel file
        from openpyxl import load_workbook
        wb = load_workbook(str(output_path))
        assert "Tabela 1" in wb.sheetnames
        ws = wb["Tabela 1"]
        assert ws.cell(1, 1).value == "Name"
        assert ws.cell(2, 1).value == "Item 1"
    
    def test_extract_multiple_tables(self, word_tool, sample_document_with_multiple_tables, tmp_path):
        """Test extracting multiple tables to Excel."""
        output_path = tmp_path / "output.xlsx"
        result = word_tool.extract_tables_to_excel(
            sample_document_with_multiple_tables,
            str(output_path)
        )
        
        assert "Extraídas 2 tabela(s)" in result
        
        # Verify both sheets exist
        from openpyxl import load_workbook
        wb = load_workbook(str(output_path))
        assert "Tabela 1" in wb.sheetnames
        assert "Tabela 2" in wb.sheetnames
    
    def test_extract_tables_custom_sheet_names(self, word_tool, sample_document_with_multiple_tables, tmp_path):
        """Test extracting tables with custom sheet names."""
        output_path = tmp_path / "output.xlsx"
        result = word_tool.extract_tables_to_excel(
            sample_document_with_multiple_tables,
            str(output_path),
            sheet_names=["Sales", "Costs"]
        )
        
        assert "Extraídas 2 tabela(s)" in result
        
        # Verify custom sheet names
        from openpyxl import load_workbook
        wb = load_workbook(str(output_path))
        assert "Sales" in wb.sheetnames
        assert "Costs" in wb.sheetnames
    
    def test_extract_tables_sheet_names_mismatch(self, word_tool, sample_document_with_multiple_tables, tmp_path):
        """Test extracting tables with mismatched sheet names."""
        output_path = tmp_path / "output.xlsx"
        
        with pytest.raises(ValueError, match="Número de nomes de abas"):
            word_tool.extract_tables_to_excel(
                sample_document_with_multiple_tables,
                str(output_path),
                sheet_names=["OnlyOne"]
            )
    
    def test_extract_tables_no_tables(self, word_tool, tmp_path):
        """Test extracting when no tables exist."""
        doc_path = tmp_path / "no_tables.docx"
        doc = Document()
        doc.add_paragraph("No tables here")
        doc.save(str(doc_path))
        
        output_path = tmp_path / "output.xlsx"
        
        with pytest.raises(ValueError, match="Nenhuma tabela encontrada"):
            word_tool.extract_tables_to_excel(str(doc_path), str(output_path))


# ==================== EXPORT TO TXT TESTS ====================

class TestExportToTxt:
    
    def test_export_to_txt_basic(self, word_tool, sample_formatted_document, tmp_path):
        """Test exporting document to TXT."""
        output_path = tmp_path / "output.txt"
        result = word_tool.export_to_txt(sample_formatted_document, str(output_path))
        
        assert "exportado para TXT" in result
        assert output_path.exists()
        
        # Verify content
        content = output_path.read_text(encoding='utf-8')
        assert "Main Title" in content
        assert "Section 1" in content
        assert "This is a paragraph" in content
    
    def test_export_to_txt_preserves_text(self, word_tool, tmp_path):
        """Test that TXT export preserves text content."""
        doc_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")
        doc.save(str(doc_path))
        
        output_path = tmp_path / "output.txt"
        word_tool.export_to_txt(str(doc_path), str(output_path))
        
        content = output_path.read_text(encoding='utf-8')
        assert "Line 1" in content
        assert "Line 2" in content
        assert "Line 3" in content


# ==================== EXPORT TO MARKDOWN TESTS ====================

class TestExportToMarkdown:
    
    def test_export_to_markdown_basic(self, word_tool, sample_formatted_document, tmp_path):
        """Test exporting document to Markdown."""
        output_path = tmp_path / "output.md"
        result = word_tool.export_to_markdown(sample_formatted_document, str(output_path))
        
        assert "exportado para Markdown" in result
        assert output_path.exists()
        
        # Verify Markdown formatting
        content = output_path.read_text(encoding='utf-8')
        assert "# Main Title" in content
        assert "## Section 1" in content
        assert "1. First item" in content or "- First item" in content
    
    def test_export_to_markdown_with_table(self, word_tool, sample_document_with_table, tmp_path):
        """Test exporting document with table to Markdown."""
        output_path = tmp_path / "output.md"
        word_tool.export_to_markdown(sample_document_with_table, str(output_path))
        
        content = output_path.read_text(encoding='utf-8')
        # Check for Markdown table format
        assert "|" in content
        assert "---" in content


# ==================== EXPORT TO HTML TESTS ====================

class TestExportToHtml:
    
    def test_export_to_html_basic(self, word_tool, sample_formatted_document, tmp_path):
        """Test exporting document to HTML."""
        output_path = tmp_path / "output.html"
        result = word_tool.export_to_html(sample_formatted_document, str(output_path))
        
        assert "exportado para HTML" in result
        assert output_path.exists()
        
        # Verify HTML structure
        content = output_path.read_text(encoding='utf-8')
        assert "<!DOCTYPE html>" in content
        assert "<html>" in content
        assert "<body>" in content
        assert "<h1>Main Title</h1>" in content
        assert "<h2>Section 1</h2>" in content
    
    def test_export_to_html_with_table(self, word_tool, sample_document_with_table, tmp_path):
        """Test exporting document with table to HTML."""
        output_path = tmp_path / "output.html"
        word_tool.export_to_html(sample_document_with_table, str(output_path))
        
        content = output_path.read_text(encoding='utf-8')
        # Check for HTML table
        assert "<table>" in content
        assert "<tr>" in content
        assert "<th>" in content or "<td>" in content


# ==================== EXPORT TO PDF TESTS ====================

class TestExportToPdf:
    
    def test_export_to_pdf_without_library(self, word_tool, sample_formatted_document, tmp_path):
        """Test exporting to PDF without docx2pdf installed."""
        output_path = tmp_path / "output.pdf"
        
        # This will likely raise ImportError unless docx2pdf is installed
        try:
            result = word_tool.export_to_pdf(sample_formatted_document, str(output_path))
            # If it succeeds, verify the result
            assert "exportado para PDF" in result
        except ImportError as e:
            # Expected if docx2pdf is not installed
            assert "docx2pdf não está instalado" in str(e)
    
    def test_export_to_pdf_nonexistent_file(self, word_tool, tmp_path):
        """Test exporting nonexistent file to PDF."""
        output_path = tmp_path / "output.pdf"
        
        with pytest.raises(FileNotFoundError):
            word_tool.export_to_pdf("nonexistent.docx", str(output_path))


# ==================== EDGE CASES ====================

class TestEdgeCases:
    
    def test_convert_list_nonexistent_file(self, word_tool):
        """Test converting list in nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.convert_list_to_table("nonexistent.docx", list_index=0)
    
    def test_convert_table_nonexistent_file(self, word_tool):
        """Test converting table in nonexistent file."""
        with pytest.raises(FileNotFoundError):
            word_tool.convert_table_to_list("nonexistent.docx", table_index=0)
    
    def test_extract_tables_nonexistent_file(self, word_tool, tmp_path):
        """Test extracting tables from nonexistent file."""
        output_path = tmp_path / "output.xlsx"
        
        with pytest.raises(FileNotFoundError):
            word_tool.extract_tables_to_excel("nonexistent.docx", str(output_path))
    
    def test_export_txt_nonexistent_file(self, word_tool, tmp_path):
        """Test exporting nonexistent file to TXT."""
        output_path = tmp_path / "output.txt"
        
        with pytest.raises(FileNotFoundError):
            word_tool.export_to_txt("nonexistent.docx", str(output_path))
    
    def test_export_markdown_nonexistent_file(self, word_tool, tmp_path):
        """Test exporting nonexistent file to Markdown."""
        output_path = tmp_path / "output.md"
        
        with pytest.raises(FileNotFoundError):
            word_tool.export_to_markdown("nonexistent.docx", str(output_path))
    
    def test_export_html_nonexistent_file(self, word_tool, tmp_path):
        """Test exporting nonexistent file to HTML."""
        output_path = tmp_path / "output.html"
        
        with pytest.raises(FileNotFoundError):
            word_tool.export_to_html("nonexistent.docx", str(output_path))
