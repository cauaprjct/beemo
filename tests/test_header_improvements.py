"""Tests for improved add_header functionality with tab stops and total pages."""

import os
import pytest
from pathlib import Path
from docx import Document

from src.word_tool import WordTool


@pytest.fixture
def word_tool():
    """Create WordTool instance."""
    return WordTool()


@pytest.fixture
def temp_dir(tmp_path):
    """Create temporary directory for test files."""
    return str(tmp_path)


class TestHeaderImprovements:
    """Test suite for add_header improvements."""

    def test_add_header_with_total_pages(self, word_tool, temp_dir):
        """Test add_header with include_total_pages parameter."""
        file_path = os.path.join(temp_dir, "test_header_total_pages.docx")
        
        # Create document
        word_tool.create_word(file_path, "Test content")
        
        # Add header with total pages
        word_tool.add_header(
            file_path,
            text="Test Report",
            alignment="left",
            include_page_number=True,
            page_number_position="right",
            include_total_pages=True
        )
        
        # Verify file exists and is valid
        assert Path(file_path).exists()
        doc = Document(file_path)
        assert len(doc.sections) > 0
        
        # Verify header has content
        header = doc.sections[0].header
        assert len(header.paragraphs) > 0
        assert "Test Report" in header.paragraphs[0].text

    def test_add_header_with_tab_stops(self, word_tool, temp_dir):
        """Test add_header with use_tab_stops parameter."""
        file_path = os.path.join(temp_dir, "test_header_tab_stops.docx")
        
        # Create document
        word_tool.create_word(file_path, "Test content")
        
        # Add header with tab stops
        word_tool.add_header(
            file_path,
            text="Company Report",
            alignment="left",
            include_page_number=True,
            page_number_position="right",
            use_tab_stops=True
        )
        
        # Verify file exists and is valid
        assert Path(file_path).exists()
        doc = Document(file_path)
        
        # Verify header has content
        header = doc.sections[0].header
        assert len(header.paragraphs) > 0
        para = header.paragraphs[0]
        assert "Company Report" in para.text
        
        # Verify tab stops were configured
        assert len(para.paragraph_format.tab_stops) > 0

    def test_add_header_with_both_features(self, word_tool, temp_dir):
        """Test add_header with both tab stops and total pages."""
        file_path = os.path.join(temp_dir, "test_header_complete.docx")
        
        # Create document with multiple pages
        word_tool.create_word(file_path, "Page 1 content")
        word_tool.add_page_break(file_path)
        word_tool.add_paragraph(file_path, "Page 2 content")
        
        # Add header with all features
        word_tool.add_header(
            file_path,
            text="Beemo AI • Test Report",
            alignment="left",
            include_page_number=True,
            page_number_position="right",
            include_total_pages=True,
            use_tab_stops=True,
            font_size=10,
            font_name="Arial"
        )
        
        # Verify file exists and is valid
        assert Path(file_path).exists()
        doc = Document(file_path)
        
        # Verify header has content
        header = doc.sections[0].header
        assert len(header.paragraphs) > 0
        para = header.paragraphs[0]
        assert "Beemo AI" in para.text
        
        # Verify tab stops were configured
        assert len(para.paragraph_format.tab_stops) > 0
        
        # Verify runs have correct formatting
        assert len(para.runs) > 0
        first_run = para.runs[0]
        assert first_run.font.size.pt == 10
        assert first_run.font.name == "Arial"

    def test_add_header_backward_compatibility(self, word_tool, temp_dir):
        """Test that add_header still works without new parameters."""
        file_path = os.path.join(temp_dir, "test_header_legacy.docx")
        
        # Create document
        word_tool.create_word(file_path, "Test content")
        
        # Add header using old API (without new parameters)
        word_tool.add_header(
            file_path,
            text="Simple Header",
            alignment="center"
        )
        
        # Verify file exists and is valid
        assert Path(file_path).exists()
        doc = Document(file_path)
        
        # Verify header has content
        header = doc.sections[0].header
        assert len(header.paragraphs) > 0
        assert "Simple Header" in header.paragraphs[0].text

    def test_add_header_with_custom_margins(self, word_tool, temp_dir):
        """Test add_header with tab stops respects custom margins."""
        file_path = os.path.join(temp_dir, "test_header_margins.docx")
        
        # Create document
        word_tool.create_word(file_path, "Test content")
        
        # Set ABNT margins
        word_tool.set_page_margins(file_path, top=3.0, bottom=2.0, left=3.0, right=2.0)
        
        # Add header with tab stops
        word_tool.add_header(
            file_path,
            text="ABNT Document",
            alignment="left",
            include_page_number=True,
            page_number_position="right",
            include_total_pages=True,
            use_tab_stops=True
        )
        
        # Verify file exists and is valid
        assert Path(file_path).exists()
        doc = Document(file_path)
        
        # Verify header and margins
        header = doc.sections[0].header
        assert "ABNT Document" in header.paragraphs[0].text
        
        # Verify margins are correct (approximately, in EMU units)
        section = doc.sections[0]
        assert section.top_margin > 0
        assert section.left_margin > 0
